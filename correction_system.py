# ================= correction_system.py =================
# ⚠️ هذا الملف مستقل تماماً - لا يعتمد على أي متغيرات خارجية
# ========================================================

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

class CorrectionSystem:
    """
    نظام تصحيح مستقل - يمكن إضافته كتاب جديد دون تعديل الكود الأصلي
    """
    
    def __init__(self):
        self.teachers_df = None
        self.halls_df = None
        self.assignments = []
        self.log = []
    
    def load_teachers(self, file_path):
        """رفع ملف المعلمين"""
        try:
            self.teachers_df = pd.read_excel(file_path) if file_path.endswith('.xlsx') else pd.read_csv(file_path)
            self._log("تم رفع بيانات المعلمين", "success")
            return True
        except Exception as e:
            self._log(f"خطأ في رفع الملف: {str(e)}", "error")
            return False
    
    def load_halls(self, file_path):
        """رفع ملف القاعات"""
        try:
            self.halls_df = pd.read_excel(file_path) if file_path.endswith('.xlsx') else pd.read_csv(file_path)
            self._log("تم رفع بيانات القاعات", "success")
            return True
        except Exception as e:
            self._log(f"خطأ في رفع ملف القاعات: {str(e)}", "error")
            return False
    
    def auto_assign(self, selected_hall=None, selected_subject=None, exam_name="امتحان عام"):
        """
        التوزيع التلقائي للمعلمين على القاعات
        - يمكن اختيار قاعة محددة أو تركها للتوزيع العشوائي
        - يمكن اختيار مبحث محدد أو تركه للجميع
        """
        if self.teachers_df is None or self.halls_df is None:
            self._log("يجب رفع البيانات أولاً", "error")
            return False
        
        # تصفية البيانات حسب الاختيارات
        filtered = self.teachers_df.copy()
        if selected_subject:
            filtered = filtered[filtered['subject'] == selected_subject]
        
        assignments = []
        
        for _, teacher in filtered.iterrows():
            # اختيار القاعة: محددة أو عشوائية من القائمة
            if selected_hall:
                hall = self.halls_df[self.halls_df['ZHALL'] == selected_hall]
            else:
                hall = self.halls_df[self.halls_df['ZLOC'] == teacher.get('city', '')]
                if hall.empty:
                    hall = self.halls_df.sample(n=1)
            
            if not hall.empty:
                hall_info = hall.iloc[0]
                
                assignment = {
                    'ZID': str(teacher['id']),
                    'ZNAME': teacher['name'],
                    'ZTEST': exam_name,
                    'ZHALL': hall_info['ZHALL'],
                    'ZLOC': hall_info['ZLOC'],
                    'ZWORK': teacher.get('work_location', 'غير محدد'),
                    'ZCITY': teacher.get('city', 'غير محدد'),
                    'subject': teacher['subject'],
                    'phone': teacher.get('phone', ''),
                    'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M")
                }
                assignments.append(assignment)
                self._log(f"تم تكليف: {teacher['name']} - {hall_info['ZHALL']}", "info")
        
        self.assignments = assignments
        return True
    
    def generate_word_letters(self, output_folder="تكاليف_التصحيح"):
        """تصدير كتب التكليف كملفات وورد فردية"""
        if not self.assignments:
            return False
        
        os.makedirs(output_folder, exist_ok=True)
        
        for assign in self.assignments:
            doc = Document()
            
            # تنسيق الرأس
            header = doc.add_heading('كتاب تكليف بتصحيح امتحان', level=1)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # محتوى الكتاب باستخدام المتغيرات المطلوبة
            content = f"""
            بناءً على صلاحياتنا الممنوحة، نكلف السيد/ة:

            الاسم: {assign['ZNAME']}
            رقم الهوية: {assign['ZID']}

            بتصحيح أوراق امتحان: {assign['ZTEST']}

            وذلك في:
            القاعة: {assign['ZHALL']}
            المدينة: {assign['ZLOC']}

            معلومات إضافية:
            مكان السكن: {assign['ZCITY']}
            مكان العمل السابق: {assign['ZWORK']}

            وتعتبر هذه المهمة جزءاً من الواجب الوطني، ونتمنى التوفيق للجميع.

            تاريخ الإصدار: {assign['timestamp']}
            """
            
            para = doc.add_paragraph(content)
            para_format = para.paragraph_format
            para_format.space_after = Pt(12)
            
            # حفظ الملف باسم الموظف ورقم هويته
            filename = f"{output_folder}/تكليف_{assign['ZNAME']}_{assign['ZID']}.docx"
            doc.save(filename)
        
        self._log(f"تم تصدير {len(self.assignments)} كتاب تكليف", "success")
        return True
    
    def export_assignments_excel(self, output_path="تكاليف_التصحيح.xlsx"):
        """تصدير جدول التكليفات كملف إكسل"""
        if self.assignments:
            df = pd.DataFrame(self.assignments)
            df.to_excel(output_path, index=False)
            self._log(f"تم تصدير ملف الإكسل: {output_path}", "success")
            return True
        return False
    
    def get_operations_log(self):
        """إرجاع سجل العمليات للعرض في التبويب المخصص"""
        return self.log
    
    def _log(self, message, level="info"):
        """تسجيل العمليات داخلياً"""
        entry = {
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'level': level,
            'message': message
        }
        self.log.append(entry)
