import streamlit as st
import pandas as pd
import sqlite3
from docx import Document
import io
import os
import time
from datetime import datetime
import copy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import zipfile  # ✅ للتصدير المضغوط

# =====================================
# 1. نظام تسجيل الدخول باستخدام Secrets + الصلاحيات
# =====================================
def login():
    if 'logged_in' not in st.session_state:
        st.session_state['logged_in'] = False
        st.session_state['username'] = ""
        st.session_state['user_role'] = "EDITOR"  # ✅ الافتراضي

    if not st.session_state['logged_in']:
        st.markdown("<h2 style='text-align: center;'>🔐 نظام تكليفات المكتب - دخول</h2>", unsafe_allow_html=True)
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            with st.form("login_form"):
                user = st.text_input("اسم المستخدم").lower().strip()
                pw = st.text_input("كلمة المرور", type="password").strip()
                submit = st.form_submit_button("دخول")
                
                if submit:
                    try:
                        valid_password = st.secrets[f"password_{user}"]
                        if pw == valid_password:
                            st.session_state['logged_in'] = True
                            st.session_state['username'] = user.upper()
                            # ✅ تحديد الصلاحية بناءً على اسم المستخدم
                            if user.upper() in ["AWAD", "SHOROQ"]:
                                st.session_state['user_role'] = "ADMIN"
                            elif user.upper() in ["MAJED", "HIND"]:
                                st.session_state['user_role'] = "EDITOR"
                            else:
                                st.session_state['user_role'] = "EDITOR"
                            st.rerun()
                        else:
                            st.error("❌ كلمة المرور غير صحيحة")
                    except KeyError:
                        st.error("❌ اسم المستخدم غير معرف في Secrets")
        return False
    return True

if not login():
    st.stop()

# =====================================
# 2. إعدادات الحالة والتبديل
# =====================================
if 'popover_counter' not in st.session_state:
    st.session_state.popover_counter = 0

if 'system_mode' not in st.session_state:
    st.session_state['system_mode'] = "tawjihi"

def switch_system(mode):
    st.session_state['system_mode'] = mode
    st.cache_data.clear()
    st.rerun()

def switch_to_other_assignments():
    st.session_state['system_mode'] = "other_assignments"
    st.cache_data.clear()
    st.rerun()

# ✅ دالة مساعدة للتحقق من الصلاحيات
def is_admin():
    return st.session_state.get('user_role') == "ADMIN"

def is_editor():
    return st.session_state.get('user_role') in ["ADMIN", "EDITOR"]

# 🔧 التعديل 1: إضافة الوضع الثالث والرابع
if st.session_state['system_mode'] == "tawjihi":
    DB_NAME = "data_system_v26.db"
    TEMPLATE_NAME = "template.docx"
    TEACHERS_URL = "https://docs.google.com/spreadsheets/d/e/2PACX-1vSubFlcocaWSvF7GU14hNGx1cuLJBwF5SchDxzeaNMJnSy6T_b0Hu5aDMnc-OM9u7EnNIATUui12H9L/pub?gid=264504938&single=true&output=csv"
    HALLS_URL = "https://docs.google.com/spreadsheets/d/e/2PACX-1vSubFlcocaWSvF7GU14hNGx1cuLJBwF5SchDxzeaNMJnSy6T_b0Hu5aDMnc-OM9u7EnNIATUui12H9L/pub?gid=1364805271&single=true&output=csv"
    PAGE_TITLE = "نظام التكليفات امتحان الثانوية العامة "
    LAST_SYNC_KEY = "last_sync_tawjihi"
elif st.session_state['system_mode'] == "tasheeh":
    DB_NAME = "data_tasheeh.db"
    TEMPLATE_NAME = "template_tasheeh.docx"
    TEACHERS_URL = "https://docs.google.com/spreadsheets/d/e/2PACX-1vRVP8cQV8GHlaWXETc9rGzteNwDVPg8iyyZ9zCXFq-J1_t0q4sxveFchsN5XbuTiZgJBeTpC3VBMc7k/pub?gid=0&single=true&output=csv"
    HALLS_URL = "https://docs.google.com/spreadsheets/d/e/2PACX-1vRVP8cQV8GHlaWXETc9rGzteNwDVPg8iyyZ9zCXFq-J1_t0q4sxveFchsN5XbuTiZgJBeTpC3VBMc7k/pub?gid=1885970999&single=true&output=csv"
    PAGE_TITLE = "نظام تصحيح الثانوية العامة"
    LAST_SYNC_KEY = "last_sync_tasheeh"
elif st.session_state['system_mode'] == "other_assignments":
    DB_NAME = "data_other_assignments.db"
    TEMPLATE_NAME = "template_other.docx"
    PAGE_TITLE = "نظام التكليفات الأخرى"
    LAST_SYNC_KEY = "last_sync_other"
else:
    DB_NAME = "data_tawzif.db"
    TEMPLATE_NAME = "template_tawzif.docx"
    TEACHERS_URL = "https://docs.google.com/spreadsheets/d/e/2PACX-1vTIka1g67VWzR7UKmdR6eb79WuCFaC-qTNTeNMYbjzkz_HmBR_Qwe6o5RGbPyPqiaY_y_z3k2YdbibO/pub?gid=821672282&single=true&output=csv"
    HALLS_URL = "https://docs.google.com/spreadsheets/d/e/2PACX-1vTIka1g67VWzR7UKmdR6eb79WuCFaC-qTNTeNMYbjzkz_HmBR_Qwe6o5RGbPyPqiaY_y_z3k2YdbibO/pub?gid=932943855&single=true&output=csv"
    PAGE_TITLE = "نظام التكليفات امتحان التوظيف"
    LAST_SYNC_KEY = "last_sync_tawzif"

st.set_page_config(page_title=PAGE_TITLE, layout="wide", initial_sidebar_state="collapsed")
st.markdown("""
    <style>
        .custom-header {
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            background-color: #1a1c23;
            color: white;
            text-align: center;
            padding: 15px 0;
            z-index: 999999;
            border-bottom: 2px solid #00ffcc;
            line-height: 1.5;
            direction: rtl;
            box-shadow: 0px 4px 10px rgba(0,0,0,0.5);
        }
        .stApp {
            margin-top: 80px;
        }
        header {visibility: hidden;}
    </style>
    
    <div class="custom-header">
        <div style="font-weight: bold; font-size: 1.2rem; text-align: center;">إعداد وتصميم : عوض نعمان ريده</div>
        <div style="font-size: 1rem; color: #00ffcc; text-align: center;">قسم الامتحانات - مديرية التربية والتعليم جنوب نابلس</div>
    </div>
    """, unsafe_allow_html=True)

st.markdown("""
    <style>
    .main, .stApp { direction: rtl; text-align: right; background-color: #0e1117; }
    .user-box { background-color: #1a1c23; padding: 5px 15px; border-radius: 8px; border-right: 5px solid #00ffcc; display: inline-block; float: right; }
    .counter-card { background-color: #1a1c23; padding: 10px; border-radius: 10px; text-align: center; border: 1px solid #333; margin-bottom: 5px; }
    .counter-label { color: #bbb; font-size: 0.85rem; }
    .counter-value { color: #00ffcc; font-size: 1.5rem; font-weight: bold; }
    div[data-baseweb="select"], div[data-baseweb="input"], .stMultiSelect { direction: rtl !important; text-align: right !important; }
    div[data-testid="stExpander"] { border: 1px solid #444 !important; background-color: #1a1c23 !important; }
    button[key^="btn_"] { background-color: #28a745 !important; color: white !important; }
    button[key^="del_"] { background-color: #dc3545 !important; color: white !important; }
    .stDownloadButton button { background-color: #007bff !important; color: white !important; }
    .editor-info { color: #ffc107 !important; font-size: 0.9rem; font-weight: bold; }
    [data-testid="stMetricValue"] { font-size: 1.8rem !important; color: #00ffcc !important; }
    .stat-card { flex: 1; padding: 15px; border-radius: 10px; text-align: center; min-width: 150px; border: 1px solid #333; }
    .stat-wants { border-top: 5px solid #28a745; background-color: #1a2e1f; }
    .stat-no-wants { border-top: 5px solid #dc3545; background-color: #2e1a1a; }
    .move-to-right { text-align: right !important; direction: rtl !important; display: block; width: 100%; color: white; }
    [data-testid="stSidebar"] { display: none; }
    /* ✅ تنسيق صفحة الطباعة */
    .print-preview { background: white; color: black; padding: 20px; border-radius: 5px; direction: rtl; text-align: right; }
    .print-preview table { width: 100%; border-collapse: collapse; }
    .print-preview td { padding: 8px; border: 1px solid #000; }
    </style>
    """, unsafe_allow_html=True)

conn = sqlite3.connect(DB_NAME, check_same_thread=False, timeout=30)
c = conn.cursor()

c.execute('''CREATE TABLE IF NOT EXISTS teachers 
             (id TEXT PRIMARY KEY, name TEXT, phone TEXT, school TEXT, city TEXT, 
             role TEXT, hall TEXT, hall_city TEXT, updated_by TEXT,
             preference TEXT, current_job TEXT, ability TEXT,
             relative TEXT, relative_exam TEXT)''')

# 🔧 التعديل 2: إضافة عمود subject
for col in ['relative', 'relative_exam', 'subject']:
    try:
        c.execute(f"ALTER TABLE teachers ADD COLUMN {col} TEXT DEFAULT ''")
        conn.commit()
    except:
        pass

c.execute('''CREATE TABLE IF NOT EXISTS halls (hall_name TEXT PRIMARY KEY, city TEXT)''')
c.execute('''CREATE TABLE IF NOT EXISTS logs 
             (id INTEGER PRIMARY KEY AUTOINCREMENT, user TEXT, action TEXT, details TEXT, timestamp TEXT,
              old_value TEXT, new_value TEXT)''')  # ✅ إضافة حقول التدقيق
conn.commit()
# ✅ إضافة أعمدة التدقيق المفصل إذا لم تكن موجودة
try:
    c.execute("ALTER TABLE logs ADD COLUMN old_value TEXT")
    c.execute("ALTER TABLE logs ADD COLUMN new_value TEXT")
    conn.commit()
except:
    pass  # الأعمدة موجودة مسبقاً
# إضافة عمود الملاحظات (مرة واحدة فقط)
try:
    c.execute("ALTER TABLE teachers ADD COLUMN notes TEXT DEFAULT ''")
    conn.commit()
except:
    pass  # العمود موجود مسبقاً

@st.cache_data(ttl=10)
def get_cached_teachers():
    return pd.read_sql("SELECT * FROM teachers", conn)

@st.cache_data(ttl=60)
def get_cached_halls():
    return pd.read_sql("SELECT * FROM halls", conn)

# ✅ دالة تسجيل التدقيق المفصل
def add_audit_log(action, details, old_value=None, new_value=None):
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    c.execute("INSERT INTO logs (user, action, details, timestamp, old_value, new_value) VALUES (?, ?, ?, ?, ?, ?)", 
              (st.session_state.username, action, details, now, 
               str(old_value) if old_value is not None else None,
               str(new_value) if new_value is not None else None))
    conn.commit()
    st.cache_data.clear()

# ✅ دالة تسجيل العمليات العادية (مطلوبة للتوافق مع الكود)
def add_log(action, details):
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    c.execute("INSERT INTO logs (user, action, details, timestamp) VALUES (?, ?, ?, ?)", 
              (st.session_state.username, action, details, now))
    conn.commit()
    st.cache_data.clear()
# =====================================
# 3. وظائف معالجة الملفات
# =====================================
def process_doc(doc_obj, row, h_name, h_city):
    phone_val = str(row.get('phone', ''))
    if phone_val.startswith('5') and len(phone_val) == 9:
        phone_val = '0' + phone_val
    
    h_name_final = str(h_name) if h_name and str(h_name).lower() != 'nan' else "---"
    h_city_final = str(h_city) if h_city and str(h_city).lower() != 'nan' else "---"
        
    repls = {
        'ZNAME': str(row.get('name', '')),
        'ZID': str(row.get('id', '')),
        'ZPHONE': phone_val,
        'ZJOB': str(row.get('role', '') or '---'),
        'ZHALL': h_name_final,
        'ZLOC': h_city_final,
        'ZWORK': str(row.get('school', '')),
        'ZCITY': str(row.get('city', '')),
        'ZREL': str(row.get('relative', 'لا يوجد')),
        'ZRELEXAM': str(row.get('relative_exam', '---')),
        # ✅ إضافة التاريخ
        'ZDATE': st.session_state.get('assign_date', datetime.now().strftime("%Y/%m/%d"))
    }

    for p in doc_obj.paragraphs:
        for k, v in repls.items():
            if k in p.text:
                for run in p.runs:
                    if k in run.text:
                        run.text = run.text.replace(k, v)
                        run.bold = True

    for table in doc_obj.tables:
        for r in table.rows:
            for cell in r.cells:
                for p in cell.paragraphs:
                    for k, v in repls.items():
                        if k in p.text:
                            for run in p.runs:
                                if k in run.text:
                                    run.text = run.text.replace(k, v)
                                    run.bold = True
    return doc_obj

def generate_single_doc(row):
    if not os.path.exists(TEMPLATE_NAME):
        st.error(f"❌ ملف القالب '{TEMPLATE_NAME}' غير موجود")
        return None
    doc = Document(TEMPLATE_NAME)
    doc = process_doc(doc, row, row['hall'], row['hall_city'])
    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    return bio

def generate_bulk_word(df, h_name):
    if not os.path.exists(TEMPLATE_NAME):
        st.error(f"❌ ملف القالب '{TEMPLATE_NAME}' غير موجود")
        return None
        
    final_doc = Document(TEMPLATE_NAME)
    final_doc._body.clear_content()
    rows_list = list(df.iterrows())
    
    for i, (idx, row) in enumerate(rows_list):
        temp_doc = Document(TEMPLATE_NAME)
        temp_doc = process_doc(temp_doc, row, h_name, row['hall_city'])
        
        elements = [el for el in temp_doc.element.body if not el.tag.endswith('sectPr')]
        
        while elements:
            last = elements[-1]
            if last.tag.endswith('}p'):
                text = ''.join(t.text or '' for t in last.iter(qn('w:t')))
                if not text.strip():
                    elements.pop()
                    continue
            break
        
        for element in elements:
            final_doc.element.body.append(copy.deepcopy(element))
        
        if i < len(rows_list) - 1:
            p = OxmlElement('w:p')
            r = OxmlElement('w:r')
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            r.append(br)
            p.append(r)
            final_doc.element.body.append(p)
    
    out = io.BytesIO()
    final_doc.save(out)
    out.seek(0)
    return out

# =====================================
# 4. الواجهة الرئيسية
# =====================================
header_col1, header_col2 = st.columns([4, 1])

with header_col1:
    # ✅ عرض الصلاحية بجانب اسم المستخدم
    role_badge = "👑 ADMIN" if is_admin() else "✏️ EDITOR"
    role_color = "#00ffcc" if is_admin() else "#ffc107"
    st.markdown(f"""
        <div class="user-box">
            <span style="color: #bbb;">👤 {st.session_state.username}</span> 
            <span style="color: {role_color}; font-weight: bold; margin-right: 10px;">{role_badge}</span>
        </div>
    """, unsafe_allow_html=True)
    
    st.write("") 
    btn_col1, btn_col2, btn_col3, btn_col4 = st.columns([1, 1, 1, 1])
    with btn_col1:
        if st.button("📝 الثانوية العامة", use_container_width=True, type="primary" if st.session_state.system_mode=="tawjihi" else "secondary"):
            switch_system("tawjihi")
    with btn_col2:
        if st.button("👨‍🏫 امتحان التوظيف", use_container_width=True, type="primary" if st.session_state.system_mode=="tawzif" else "secondary"):
            switch_system("tawzif")
    with btn_col3:
        if st.button("✅ تصحيح الثانوية", use_container_width=True, type="primary" if st.session_state.system_mode=="tasheeh" else "secondary"):
            switch_system("tasheeh")
    with btn_col4:
        if st.button("📋 تكليفات أخرى", use_container_width=True, type="primary" if st.session_state.system_mode=="other_assignments" else "secondary"):
            switch_to_other_assignments()

with header_col2:
    if st.button("🚪 تسجيل الخروج", key="logout_btn", use_container_width=True):
        st.session_state.logged_in = False
        st.rerun()

st.divider()

# 🔴🔴🔴 التبويبات الأصلية تظهر فقط إذا لم يكن الوضع "تصحيح الثانوية" 🔴🔴🔴
if st.session_state['system_mode'] not in ["tasheeh", "other_assignments"]:
    
    tab_search, tab_auto, tab_upload, tab_manage, tab_logs, tab_inquiry = st.tabs([
    "🔍 البحث والتعيين", "🤖 التوزيع التلقائي", "📥 رفع البيانات", 
    "📊 الإدارة والإحصائيات", "📜 سجل العمليات", "🔎 الاستعلامات الذكية"
])

    # ==================== تبويب البحث ====================
    with tab_search:
        st.markdown(f'<h2 class="move-to-right">إدارة الموظفين - {PAGE_TITLE}</h2>', unsafe_allow_html=True)
        
        # 📅 حقل تاريخ التكليف (للتوزيع الفردي)
        if 'assign_date' not in st.session_state:
            st.session_state.assign_date = datetime.now().strftime("%Y/%m/%d")
        col_date1, col_date2 = st.columns([4, 1])
        with col_date1:
            st.session_state.assign_date = st.date_input(
                "📅 تاريخ التكليف:", 
                value=datetime.strptime(st.session_state.assign_date, "%Y/%m/%d"),
                key="assign_date_search"
            ).strftime("%Y/%m/%d")
        with col_date2:
            st.info(f"📌 `{st.session_state.assign_date}`")
        st.divider()
        
        df_h_data = get_cached_halls()
        hall_map = {r['hall_name']: r['city'] for _, r in df_h_data.iterrows()}
        
        # ✅ البحث الفوري (يظهر النتائج أثناء الكتابة)
        q = st.text_input("🔍 ابحث عن الاسم، الهوية، أو الجوال", key="search_live")
        
        if q:
            df_teachers = get_cached_teachers()
            results = df_teachers[df_teachers['name'].str.contains(q, na=False, case=False) | df_teachers['id'].astype(str).str.contains(q) | df_teachers['phone'].astype(str).str.contains(q)]
            
            for idx, row in results.iterrows():
                display_phone = str(row.get('phone', '---'))
                if display_phone.startswith('5') and len(display_phone) == 9:
                    display_phone = '0' + display_phone

                with st.expander(f"👤 {row.get('name', 'اسم غير معروف')} | القاعة: {row.get('hall') or 'غير مكلف'}"):
                    def safe_get(key):
                        v = str(row.get(key, '---')).strip()
                        return '---' if v.lower() in ['nan', 'none', ''] else v

                    v_id = safe_get('id')
                    v_city = safe_get('city')
                    v_school = safe_get('school')
                    v_job = safe_get('current_job')
                    v_abil = safe_get('ability')
                    v_notes = safe_get('notes')
                    v_pref = safe_get('preference')

                    rel_html = ''
                    if st.session_state.system_mode == 'tawzif':
                        v_rel = safe_get('relative')
                        v_relex = safe_get('relative_exam')
                        rel_html = f'<tr><td style="padding: 5px; color: #ffc107;"><b>🔗 قريب:</b> {v_rel}</td><td style="padding: 5px; color: #ffc107;"><b>📝 الامتحان:</b> {v_relex}</td></tr>'

                    full_table = f'<div style="background-color: #1a1c23; padding: 15px; border-radius: 10px; border: 1px solid #444; border-right: 5px solid #00ffcc; margin-bottom: 15px; text-align: right; direction: rtl;"><table style="width:100%; color: white; border: none;"><tr><td style="padding: 5px;"><b>🆔 الهوية:</b> {v_id}</td><td style="padding: 5px;"><b>📱 الجوال:</b> {display_phone}</td></tr><tr><td style="padding: 5px;"><b>🏡 السكن:</b> {v_city}</td><td style="padding: 5px;"><b>🏫 المدرسة:</b> {v_school}</td></tr><tr><td style="padding: 5px;"><b>📝 الرغبة:</b> {v_pref}</td><td style="padding: 5px;"><b>💼 الوظيفة:</b> {v_job}</td></tr>{rel_html}<tr><td colspan="2" style="padding: 5px; border-top: 1px solid #444; color: #ffc107;"><b>⚠️ صلاحية المراقبة:</b> {v_abil}</td></tr><tr><td colspan="2" style="padding: 5px; border-top: 1px solid #444; color: #888;"><b>📝 ملاحظات:</b> {v_notes if v_notes != "---" else "لا يوجد"}</td></tr></table></div>'
                    st.markdown(full_table, unsafe_allow_html=True)
                    st.markdown(f"<span class='editor-info'>آخر تعديل: {row['updated_by'] or 'لا يوجد'}</span>", unsafe_allow_html=True)
                    
                    with st.popover("📝 تعديل البيانات", key=f"pop_{row['id']}_{idx}_{st.session_state.popover_counter}"):
                        u_name = st.text_input("الاسم", value=row['name'], key=f"un_{st.session_state.system_mode}_{row['id']}_{idx}")
                        u_phone = st.text_input("رقم الجوال", value=display_phone, key=f"up_{st.session_state.system_mode}_{row['id']}_{idx}")
                        u_school = st.text_input("المدرسة", value=row['school'], key=f"us_{st.session_state.system_mode}_{row['id']}_{idx}")
                        u_city = st.text_input("السكن", value=row['city'], key=f"uc_{st.session_state.system_mode}_{row['id']}_{idx}")
                        u_job = st.text_input("الوظيفة الأساسية", value=row['current_job'], key=f"uj_{st.session_state.system_mode}_{row['id']}_{idx}")
                        
                        u_pref = st.selectbox("الرغبة", ["يرغب", "لا يرغب", "غير محدد"], 
                                            index=0 if row['preference']=="يرغب" else (1 if row['preference']=="لا يرغب" else 2), 
                                            key=f"upr_{st.session_state.system_mode}_{row['id']}_{idx}")
                        
                        u_abil = st.selectbox("صلاحية المراقبة", ["يصلح", "لا يصلح", "لم تحدد"], 
                                            index=0 if row['ability']=="يصلح" else (1 if row['ability']=="لا يصلح" else 2), 
                                            key=f"uab_{st.session_state.system_mode}_{row['id']}_{idx}")
                        u_notes = st.text_area("📝 ملاحظات إضافية", value=row.get('notes', ''), height=80, key=f"unotes_{st.session_state.system_mode}_{row['id']}_{idx}")

                        if st.session_state.system_mode == "tawzif":
                            u_rel = st.selectbox("هل له قريب؟", ["نعم", "لا"], index=0 if row.get('relative')=="نعم" else 1, key=f"urel_{row['id']}_{idx}")
                            u_relex = st.text_input("اسم امتحان القريب", value=row.get('relative_exam', ''), key=f"urex_{row['id']}_{idx}")

                        if st.button("💾 تحديث وحفظ", key=f"save_base_{row['id']}_{idx}_{st.session_state.popover_counter}"):
                            if st.session_state.system_mode == "tawzif":
                                # ✅ تسجيل التدقيق المفصل
                                old_data = row.to_dict()
                                c.execute("""UPDATE teachers SET name=?, phone=?, school=?, city=?, current_job=?, preference=?, ability=?, relative=?, relative_exam=?, notes=?, updated_by=? WHERE id=?""", 
                                         (u_name, u_phone, u_school, u_city, u_job, u_pref, u_abil, u_rel, u_relex, u_notes, st.session_state.username, row['id']))
                                new_data = {'name': u_name, 'phone': u_phone, 'school': u_school, 'city': u_city, 'current_job': u_job, 'preference': u_pref, 'ability': u_abil, 'relative': u_rel, 'relative_exam': u_relex, 'notes': u_notes}
                                add_audit_log("تعديل بيانات أساسية", f"تعديل بيانات {u_name}", old_data, new_data)
                            else:
                                old_data = row.to_dict()
                                c.execute("""UPDATE teachers SET name=?, phone=?, school=?, city=?, current_job=?, preference=?, ability=?, notes=?, updated_by=? WHERE id=?""", 
                                         (u_name, u_phone, u_school, u_city, u_job, u_pref, u_abil, u_notes, st.session_state.username, row['id']))
                                new_data = {'name': u_name, 'phone': u_phone, 'school': u_school, 'city': u_city, 'current_job': u_job, 'preference': u_pref, 'ability': u_abil, 'notes': u_notes}
                                add_audit_log("تعديل بيانات أساسية", f"تعديل بيانات {u_name}", old_data, new_data)
                            conn.commit()
                            st.session_state.popover_counter += 1
                            st.cache_data.clear()
                            st.success("✅ تم الحفظ")
                            time.sleep(0.5)
                            st.rerun()

                    st.divider()
                    c1, c2 = st.columns(2)
                    with c1:
                        current_hall = row['hall'] if row['hall'] and str(row['hall']).lower() != 'nan' else ""
                        sel_h = st.selectbox("القاعة", [""] + list(hall_map.keys()),
                                            index=(list(hall_map.keys()).index(current_hall)+1 if current_hall in hall_map else 0),
                                            key=f"q_h_{st.session_state.system_mode}_{row['id']}_{idx}")
                        sel_r = st.selectbox("المهمة", ["", "رئيس قاعة", "مساعد رئيس قاعة", "مراقب", "آذن"],
                                            index=(["", "رئيس قاعة", "مساعد رئيس قاعة", "مراقب", "آذن"].index(row['role']) if row['role'] in ["", "رئيس قاعة", "مساعد رئيس قاعة", "مراقب", "آذن"] else 0),
                                            key=f"q_r_{st.session_state.system_mode}_{row['id']}_{idx}")
                    with c2:
                        if st.button("💾 حفظ التكليف", key=f"btn_save_{st.session_state.system_mode}_{row['id']}_{idx}"):
                            if row['preference'] == 'لا يرغب':
                                st.error("⚠️ هذا المعلم لا يرغب في التكليف، يرجى تغيير حالته أولاً")
                            elif row['ability'] == 'لا يصلح':
                                st.error("⚠️ هذا المعلم لا يصلح للمراقبة، يرجى تغيير حالته أولاً")
                            else:
                                h_city_val = hall_map.get(sel_h, "")
                                old_hall = row['hall']
                                c.execute("UPDATE teachers SET hall=?, role=?, hall_city=?, updated_by=? WHERE id=?", 
                                          (sel_h, sel_r, h_city_val, st.session_state.username, row['id']))
                                conn.commit()
                                add_audit_log("حفظ تكليف", f"تم تكليف {row['name']} في {sel_h}", old_hall, sel_h)
                                st.success("✅ تم الحفظ")
                                time.sleep(0.5)
                                st.rerun()
                        
                        is_assigned = row['hall'] and str(row['hall']).strip() != "" and str(row['hall']).lower() != 'nan'
                        if is_assigned:
                            if st.button("❌ إلغاء التكليف", key=f"del_search_{st.session_state.system_mode}_{row['id']}"):
                                old_hall = row['hall']
                                c.execute("UPDATE teachers SET hall='', role='', hall_city='', updated_by=? WHERE id=?", 
                                          (st.session_state.username, row['id']))
                                conn.commit()
                                add_audit_log("إلغاء تكليف", f"تم إلغاء تكليف {row['name']}", old_hall, "")
                                st.rerun()
                            
                            if st.button("📥 إنشاء الكتاب", key=f"gen_s_{st.session_state.system_mode}_{row['id']}"):
                                f_word = generate_single_doc(row)
                                if f_word: 
                                    st.download_button("📥 تحميل الآن", data=f_word, file_name=f"تكليف_{row['name']}.docx", key=f"dl_s_{st.session_state.system_mode}_{row['id']}")
                            

    # ==================== تبويب التوزيع التلقائي ====================
        # ==================== تبويب التوزيع التلقائي ====================
    with tab_auto:
        st.markdown('<h2 class="move-to-right">🤖 نظام التوزيع التلقائي الذكي</h2>', unsafe_allow_html=True)
        
        # 📅 حقل تاريخ التكليف (للتوزيع الجماعي)
        if 'assign_date_bulk' not in st.session_state:
            st.session_state.assign_date_bulk = datetime.now().strftime("%Y/%m/%d")
        col_date_bulk1, col_date_bulk2 = st.columns([4, 1])
        with col_date_bulk1:
            st.session_state.assign_date_bulk = st.date_input(
                "📅 تاريخ التكليف الجماعي:", 
                value=datetime.strptime(st.session_state.assign_date_bulk, "%Y/%m/%d"),
                key="assign_date_bulk_auto"
            ).strftime("%Y/%m/%d")
        with col_date_bulk2:
            st.info(f"📌 `{st.session_state.assign_date_bulk}`")
        st.divider()
        
        # ✅ تعريف المتغيرات المطلوبة (بدون عرض لوحة التحكم)
        df_all = get_cached_teachers()  # ← هذا السطر كان مفقوداً!
        hall_map_auto = {r['hall_name']: r['city'] for _, r in get_cached_halls().iterrows()}
        
        df_qualified = df_all[(df_all['ability'] == 'يصلح') & (df_all['preference'] == 'يرغب') & (df_all['current_job'] == 'معلم') & ((df_all['hall'] == '') | (df_all['hall'].isna()))]
        can_and_wants = len(df_qualified)
        can_not_wants = len(df_all[(df_all['ability'] == 'يصلح') & (df_all['preference'] == 'لا يرغب') & (df_all['current_job'] == 'معلم') & ((df_all['hall'] == '') | (df_all['hall'].isna()))])
        
        st.markdown(f"""
        <div style="display: flex; gap: 15px; margin-bottom: 20px; direction: rtl;">
            <div class="stat-card stat-wants"><span style="color: #bbb; font-size: 0.9rem;">متاح (يصلح ويرغب)</span><br><strong style="font-size: 2rem; color: #28a745;">{can_and_wants}</strong></div>
            <div class="stat-card stat-no-wants"><span style="color: #bbb; font-size: 0.9rem;">متاح (يصلح ولا يرغب)</span><br><strong style="font-size: 2rem; color: #dc3545;">{can_not_wants}</strong></div>
        </div>
        """, unsafe_allow_html=True)

        available_cities = sorted(df_qualified['city'].unique().tolist()) if not df_qualified.empty else []
        col_a1, col_a2 = st.columns(2)
        with col_a1:
            target_h = st.selectbox("اختر القاعة المستهدفة:", [""] + list(hall_map_auto.keys()), key="auto_target_h")
            selected_cities = st.multiselect("السحب من مناطق سكن محددة (اختياري):", available_cities)
        with col_a2:
            df_pool = df_qualified[df_qualified['city'].isin(selected_cities)] if selected_cities else df_qualified
            st.info(f"عدد المعلمين المتاحين للسحب الآن: {len(df_pool)}")
            num_to_assign = st.number_input("العدد المطلوب توزيعه:", min_value=0, max_value=len(df_pool) if not df_pool.empty else 0, value=0)

            if st.button("🚀 ابدأ التوزيع التلقائي الآن", use_container_width=True, disabled=(num_to_assign == 0 or not target_h)):
                # 🔴🔴 الحصول على معلومات القاعة المستهدفة 🔴🔴
                target_hall_city = hall_map_auto.get(target_h, "")
                
                # 🔴 فلترة المعلمين المستبعدين (للثانوية العامة فقط)
                if st.session_state.system_mode == "tawjihi":
                    # استبعاد من يسكن في نفس مدينة القاعة
                    df_pool_filtered = df_pool[df_pool['city'] != target_hall_city].copy()
                    
                    # استبعاد من يعمل في نفس المدرسة/القاعة
                    df_pool_filtered = df_pool_filtered[df_pool_filtered['school'] != target_h].copy()
                    
                    excluded_count = len(df_pool) - len(df_pool_filtered)
                    if excluded_count > 0:
                        st.info(f"ℹ️ تم استبعاد `{excluded_count}` معلم (من نفس المدينة أو المدرسة)")
                    
                    df_pool = df_pool_filtered
                
                # التأكد من وجود عدد كافٍ للتوزيع
                actual_num = min(int(num_to_assign), len(df_pool))
                
                if actual_num == 0:
                    st.warning("⚠️ لا يوجد معلمين متاحين للتوزيع بعد تطبيق الشروط")
                else:
                    selected_sample = df_pool.sample(n=actual_num)
                    for _, r in selected_sample.iterrows():
                        old_hall = r['hall']
                        c.execute("UPDATE teachers SET hall=?, role='مراقب', hall_city=?, updated_by='توزيع تلقائي' WHERE id=?", 
                                  (target_h, target_hall_city, r['id']))
                        add_audit_log("توزيع تلقائي", f"توزيع {r['name']} على قاعة {target_h}", old_hall, target_h)
                    conn.commit()
                    add_log("توزيع تلقائي", f"توزيع {actual_num} معلم على قاعة {target_h}")
                    st.success(f"✅ تم توزيع {actual_num} بنجاح!")
                    time.sleep(1)
                    st.rerun()

        st.divider()
        st.markdown('<h3 class="move-to-right">👔 تعيين رئيس القاعة والمساعد والآذن</h3>', unsafe_allow_html=True)

        df_managers = df_all[(df_all['current_job'] == 'مدير مدرسة') & (df_all['preference'] == 'يرغب') & ((df_all['hall'].isna()) | (df_all['hall'].astype(str).str.strip().isin(['', 'nan', 'None', 'NaN'])))]
        df_secretaries = df_all[(df_all['current_job'] == 'سكرتير') & (df_all['preference'] == 'يرغب') & ((df_all['hall'].isna()) | (df_all['hall'].astype(str).str.strip().isin(['', 'nan', 'None', 'NaN'])))]
        df_janitors = df_all[(df_all['current_job'] == 'آذن') & (df_all['preference'] == 'يرغب') & ((df_all['hall'].isna()) | (df_all['hall'].astype(str).str.strip().isin(['', 'nan', 'None', 'NaN'])))]

        target_h2 = ""
        col_r1, col_r2 = st.columns(2)
        with col_r1:
            target_h2 = st.selectbox("اختر القاعة:", [""] + list(hall_map_auto.keys()), key="role_target_h")
        with col_r2:
            st.info(f"مدراء متاحين: {len(df_managers)} | سكرتارية: {len(df_secretaries)} | آذنة: {len(df_janitors)}")

        if target_h2:
            col_s1, col_s2, col_s3 = st.columns(3)
            
            with col_s1:
                sel_manager = st.selectbox("👑 رئيس القاعة (مدير مدرسة):", [""] + df_managers['name'].tolist(), key="sel_manager")
            with col_s2:
                # ✅ التعديل هنا: قائمة متعددة تتيح اختيار مساعد أو مساعدان
                sel_secretaries = st.multiselect("📋 مساعدي الرئيس (بحد أقصى 2):", df_secretaries['name'].tolist(), max_selections=2, key="sel_secretaries_multi")
            with col_s3:
                sel_janitor = st.selectbox("🔑 الآذن:", [""] + df_janitors['name'].tolist(), key="sel_janitor")
            
            if st.button("💾 حفظ التعيينات", use_container_width=True, key="save_roles"):
                saved = []
                
                # حفظ رئيس القاعة
                if sel_manager:
                    manager_id = df_managers[df_managers['name'] == sel_manager]['id'].values[0]
                    old_hall = df_managers[df_managers['name'] == sel_manager]['hall'].values[0]
                    c.execute("UPDATE teachers SET hall=?, role='رئيس قاعة', hall_city=?, updated_by=? WHERE id=?",
                              (target_h2, hall_map_auto[target_h2], st.session_state.username, manager_id))
                    add_audit_log("تعيين رئيس قاعة", f"تعيين {sel_manager}", old_hall, target_h2)
                    saved.append(f"رئيس قاعة: {sel_manager}")
                
                # ✅ حفظ المساعدين (واحد أو اثنين)
                for sec_name in sel_secretaries:
                    if sec_name:
                        sec_id = df_secretaries[df_secretaries['name'] == sec_name]['id'].values[0]
                        old_hall = df_secretaries[df_secretaries['name'] == sec_name]['hall'].values[0]
                        c.execute("UPDATE teachers SET hall=?, role='مساعد رئيس قاعة', hall_city=?, updated_by=? WHERE id=?",
                                  (target_h2, hall_map_auto[target_h2], st.session_state.username, sec_id))
                        add_audit_log("تعيين مساعد رئيس", f"تعيين {sec_name}", old_hall, target_h2)
                        saved.append(f"مساعد رئيس: {sec_name}")
                
                # حفظ الآذن
                if sel_janitor:
                    janitor_id = df_janitors[df_janitors['name'] == sel_janitor]['id'].values[0]
                    old_hall = df_janitors[df_janitors['name'] == sel_janitor]['hall'].values[0]
                    c.execute("UPDATE teachers SET hall=?, role='آذن', hall_city=?, updated_by=? WHERE id=?",
                              (target_h2, hall_map_auto[target_h2], st.session_state.username, janitor_id))
                    add_audit_log("تعيين آذن", f"تعيين {sel_janitor}", old_hall, target_h2)
                    saved.append(f"آذن: {sel_janitor}")
                
                if saved:
                    conn.commit()
                    add_log("تعيين أدوار", f"قاعة {target_h2}: {' | '.join(saved)}")
                    st.success(f"✅ تم الحفظ: {' | '.join(saved)}")
                    time.sleep(0.5)
                    st.cache_data.clear()
                    st.rerun()
                else:
                    st.warning("⚠️ لم تختر أي شخص!")
    # ==================== تبويب رفع البيانات ====================
    with tab_upload:
        st.markdown(f'<h2 class="move-to-right">تحديث البيانات - {PAGE_TITLE}</h2>', unsafe_allow_html=True)
        
        # ✅ تم حذف رفع القالب لأنه أصبح على جيت هب
       
        
        # ✅ التحقق من المزامنة التلقائية (رقم 6)
        last_sync = st.session_state.get(LAST_SYNC_KEY, "لم تتم المزامنة بعد")
        st.caption(f"🕐 آخر مزامنة: {last_sync}")
        
        st.divider()
        if st.button("🗑️ مسح البيانات المكررة"):
            try:
                c.execute("DELETE FROM teachers WHERE rowid NOT IN (SELECT MIN(rowid) FROM teachers GROUP BY id)")
                conn.commit()
                st.cache_data.clear()
                st.success("✅ تم مسح التكرار بنجاح")
                st.rerun()
            except Exception as e:
                st.error(f"خطأ: {e}")

        if st.button("🔄 تحديث من Google Sheets"):
            try:
                conn.execute("PRAGMA journal_mode=WAL")
                conn.commit()
                dft = pd.read_csv(TEACHERS_URL, dtype={'id': str, 'phone': str})
                dft.columns = dft.columns.str.strip().str.lower()
                if 'id_number' in dft.columns:
                    dft.rename(columns={'id_number': 'id'}, inplace=True)
                dft.to_sql('teachers_temp', conn, if_exists='replace', index=False)
                ids_in_sheet = dft['id'].astype(str).tolist()
                placeholders = ','.join(['?' for _ in ids_in_sheet])
                c.execute(f"DELETE FROM teachers WHERE id NOT IN ({placeholders})", ids_in_sheet)
                conn.commit()
                
                if st.session_state['system_mode'] == 'tawjihi':
                    c.execute("UPDATE teachers SET name = t.name, phone = t.phone, school = t.school, city = t.city, current_job = t.current_job, preference = t.preference, ability = t.ability FROM teachers_temp t WHERE teachers.id = t.id")
                    c.execute("INSERT OR IGNORE INTO teachers (id, name, phone, school, city, current_job, preference, ability) SELECT id, name, phone, school, city, current_job, preference, ability FROM teachers_temp")
                else:
                    c.execute("UPDATE teachers SET name = t.name, phone = t.phone, school = t.school, city = t.city, current_job = t.current_job, preference = t.preference, ability = t.ability, relative = t.relative, relative_exam = t.relative_exam FROM teachers_temp t WHERE teachers.id = t.id")
                    c.execute("INSERT OR IGNORE INTO teachers (id, name, phone, school, city, current_job, preference, ability, relative, relative_exam) SELECT id, name, phone, school, city, current_job, preference, ability, relative, relative_exam FROM teachers_temp")
                conn.commit()
                dfh = pd.read_csv(HALLS_URL)
                dfh.to_sql('halls', conn, if_exists='replace', index=False)
                conn.commit()
                
                # ✅ تحديث وقت المزامنة
                st.session_state[LAST_SYNC_KEY] = datetime.now().strftime("%Y-%m-%d %H:%M")
                
                add_log("تحديث بيانات", "تحديث ذكي من جوجل شيت (حفظ التكليفات)")
                st.success("✅ تم التحديث بنجاح مع الحفاظ على التكليفات الحالية")
                st.cache_data.clear()
                st.rerun()
            except Exception as e:
                st.error(f"خطأ أثناء التحديث: {e}")

    # ==================== تبويب الإدارة ====================
        # ==================== تبويب الإدارة ====================
    with tab_manage:
        # ✅ لوحة التحكم الإحصائية المتقدمة
        st.markdown("### 📊 لوحة التحكم المتقدمة")
        df_all_teachers = get_cached_teachers()
        
        total_count = len(df_all_teachers[(df_all_teachers['ability'] == 'يصلح') & (df_all_teachers['preference'] == 'يرغب') & (df_all_teachers['current_job'] == 'معلم')])
        assigned_count = len(df_all_teachers[(df_all_teachers['ability'] == 'يصلح') & (df_all_teachers['preference'] == 'يرغب') & (df_all_teachers['current_job'] == 'معلم') & (df_all_teachers['hall'].astype(str).str.len() > 0)])
        remaining_count = total_count - assigned_count
        
        c_m1, c_m2, c_m3 = st.columns(3)
        c_m1.metric("إجمالي الموظفين المتاحين للمراقبة", total_count)
        c_m2.metric("تم إنجازهم", assigned_count)
        c_m3.metric("المتبقي", remaining_count)
        
       
        st.divider()
        st.markdown('<h3 class="move-to-right">📦 تصدير البيانات المعدلة</h3>', unsafe_allow_html=True)
        df_export = df_all_teachers.copy()
        original_order = ['id', 'name', 'phone', 'school', 'city', 'role', 'hall', 'hall_city', 'preference', 'modified_by', 'job_title', 'permissions', 'relative', 'relative_exam']
        existing_cols = [c for c in original_order if c in df_export.columns]
        df_final = df_export[existing_cols].copy()
        column_mapping = {'id': 'رقم الهوية', 'name': 'الاسم كامل', 'phone': 'رقم الجوال', 'school': 'المدرسة', 'city': 'السكن', 'role': 'المهمة المكلف بها', 'hall': 'القاعة', 'hall_city': 'مدينة القاعة', 'preference': 'الرغبة', 'modified_by': 'الموظف المعدل', 'job_title': 'الوظيفة', 'permissions': 'الصلاحية', 'relative': 'قريب مباشر', 'relative_exam': 'امتحان القريب'}
        df_final.rename(columns=column_mapping, inplace=True)

        output_all = io.BytesIO()
        with pd.ExcelWriter(output_all, engine='xlsxwriter') as writer:
            df_final.to_excel(writer, index=False, sheet_name='الموظفين')
            workbook = writer.book
            worksheet = writer.sheets['الموظفين']
            h_fmt = workbook.add_format({'bold': True, 'font_size': 14, 'border': 1, 'align': 'center', 'bg_color': '#D7E4BC'})
            c_fmt = workbook.add_format({'bold': True, 'font_size': 14, 'border': 1, 'align': 'right'})
            worksheet.right_to_left()
            worksheet.set_landscape()
            worksheet.fit_to_pages(1, 0)
            for col_num, col_name in enumerate(df_final.columns):
                worksheet.write(0, col_num, col_name, h_fmt)
                column_data = df_final[col_name].astype(str).str.len()
                max_len = max(column_data.max() if not column_data.empty else 0, len(str(col_name))) + 7
                worksheet.set_column(col_num, col_num, min(max_len, 50), c_fmt)

        st.download_button(label="📥 تحميل إكسل معدل", data=output_all.getvalue(), file_name=f"كشف_عام_{datetime.now().strftime('%Y%m%d')}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        st.divider()
        st.markdown("### 📦 التصدير الجماعي")
        assigned_halls = sorted(df_all_teachers[df_all_teachers['hall'].astype(str).str.len() > 0]['hall'].unique().tolist())
        
        if assigned_halls:
            if st.button("📦 إنشاء ملف مضغوط لجميع التكليفات", type="primary", use_container_width=True):
                with st.spinner("جاري إنشاء الملف المضغوط..."):
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                        for hall in assigned_halls:
                            df_hall = df_all_teachers[df_all_teachers['hall'] == hall]
                            for _, row in df_hall.iterrows():
                                # ✅ التعديل هنا: الدالة تعيد الملف جاهزاً، نأخذ المحتوى مباشرة
                                doc_buffer = generate_single_doc(row)
                                if doc_buffer:
                                    filename = f"تكليف_{row['name']}_{row['id']}.docx"
                                    zip_file.writestr(filename, doc_buffer.getvalue())
                    
                    zip_buffer.seek(0)
                    st.download_button(
                        label="📥 تحميل الملف المضغوط",
                        data=zip_buffer.getvalue(),
                        file_name=f"تكليفات_جميع_القاعات_{datetime.now().strftime('%Y%m%d')}.zip",
                        mime="application/zip"
                    )
                    st.success("✅ تم إنشاء الملف المضغوط بنجاح!")
        
        # ✅ عرض تفاصيل القاعة (مع تعريف الأعمدة بشكل آمن)
        if assigned_halls:
            h_choice = st.selectbox("اختر قاعة لعرض الكادر والإحصائيات:", [""] + assigned_halls)
            if h_choice:
                df_hall_details = df_all_teachers[df_all_teachers['hall'] == h_choice].copy()
                st.markdown(f'<h4 class="move-to-right">🔢 معداد قاعة: {h_choice}</h4>', unsafe_allow_html=True)
                m1, m2, m3, m4 = st.columns(4)
                with m1: st.markdown(f'<div class="counter-card"><div class="counter-label">رئيس قاعة</div><div class="counter-value">{len(df_hall_details[df_hall_details["role"] == "رئيس قاعة"])}</div></div>', unsafe_allow_html=True)
                with m2: st.markdown(f'<div class="counter-card"><div class="counter-label">مساعد رئيس</div><div class="counter-value">{len(df_hall_details[df_hall_details["role"] == "مساعد رئيس قاعة"])}</div></div>', unsafe_allow_html=True)
                with m3: st.markdown(f'<div class="counter-card"><div class="counter-label">مراقب</div><div class="counter-value">{len(df_hall_details[df_hall_details["role"] == "مراقب"])}</div></div>', unsafe_allow_html=True)
                with m4: st.markdown(f'<div class="counter-card"><div class="counter-label">آذن</div><div class="counter-value">{len(df_hall_details[df_hall_details["role"] == "آذن"])}</div></div>', unsafe_allow_html=True)
                
                st.divider()
                if not df_hall_details.empty:
                    df_to_show = df_hall_details[['name', 'role', 'school', 'city', 'phone']].copy()
                    df_to_show.insert(0, 'م', range(1, 1 + len(df_to_show)))
                    df_to_show.columns = ['الرقم', 'الاسم', 'المهمة', 'المدرسة', 'السكن', 'الجوال']
                    styled_df = df_to_show.style.set_properties(**{'text-align': 'right', 'direction': 'rtl'}).hide(axis="index")
                    st.markdown(styled_df.to_html(), unsafe_allow_html=True)

                st.markdown("<br>", unsafe_allow_html=True)
                
                # ✅ تعريف الأعمدة هنا لضمان وجودها قبل الاستخدام
                col_btns1, col_btns2, col_btns3 = st.columns([1, 1.2, 1.2])
                
                with col_btns1:
                    if st.button(f"🗑️ تفريغ قاعة {h_choice}", key=f"del_hall_{h_choice}"):
                        c.execute("UPDATE teachers SET hall='', role='', hall_city='', updated_by=? WHERE hall=?", (st.session_state.username, h_choice))
                        conn.commit()
                        add_log("تفريغ قاعة", f"تم مسح كافة تكليفات قاعة {h_choice}")
                        st.success("تم تفريغ القاعة")
                        time.sleep(0.5)
                        st.rerun()
                with col_btns2:
                    if st.button(f"📄 إنشاء كتب قاعة {h_choice}", key=f"gen_bulk_{h_choice}"):
                        bulk_f = generate_bulk_word(df_hall_details, h_choice)
                        if bulk_f: 
                            st.download_button("📥 تحميل الوورد", data=bulk_f, file_name=f"تكليفات_{h_choice}.docx")
                with col_btns3:
                    output_hall_excel = io.BytesIO()
                    df_hall_excel = df_hall_details.copy()
                    df_hall_excel.insert(0, 'الرقم', range(1, 1 + len(df_hall_excel)))
                    df_final_export = df_hall_excel[['الرقم', 'name', 'id', 'phone', 'school', 'role', 'city']]
                    df_final_export.columns = ['الرقم', 'الاسم الرباعي', 'رقم الهوية', 'رقم الجوال', 'المدرسة', 'المهمة', 'العنوان']
                    with pd.ExcelWriter(output_hall_excel, engine='xlsxwriter') as writer:
                        df_final_export.to_excel(writer, index=False, sheet_name='كشف_القاعة', startrow=1)
                        workbook = writer.book
                        worksheet = writer.sheets['كشف_القاعة']
                        title_fmt = workbook.add_format({'bold': True, 'font_size': 20, 'border': 1, 'align': 'center', 'valign': 'vcenter', 'bg_color': '#BDD7EE'})
                        h_fmt = workbook.add_format({'bold': True, 'font_size': 14, 'border': 1, 'align': 'center', 'valign': 'vcenter', 'bg_color': '#BDD7EE'})
                        c_fmt = workbook.add_format({'bold': True, 'font_size': 14, 'border': 1, 'align': 'right', 'valign': 'vcenter'})
                        worksheet.right_to_left()
                        worksheet.set_landscape()
                        worksheet.fit_to_pages(1, 0)
                        header_text = f"بيانات قاعة: {h_choice}"
                        worksheet.merge_range(0, 0, 0, 6, header_text, title_fmt)
                        worksheet.set_row(0, 35)
                        for col_num, col_name in enumerate(df_final_export.columns):
                            worksheet.write(1, col_num, col_name, h_fmt)
                            column_length = max(df_final_export[col_name].astype(str).map(len).max(), len(str(col_name))) + 4
                            worksheet.set_column(col_num, col_num, column_length, c_fmt)
                    add_log("تصدير إكسل", f"تحميل كشف قاعة: {h_choice}")
                    st.download_button(label=f"📊 كشف إكسل {h_choice}", data=output_hall_excel.getvalue(), file_name=f"كشف_{h_choice}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", key=f"dl_xl_{h_choice}_export")

    # ==================== تبويب السجلات ====================
    with tab_logs:
        st.markdown('<h2 class="move-to-right">📜 سجل العمليات</h2>', unsafe_allow_html=True)
        if is_admin():  # ✅ فقط الأدمن يمكنه حذف السجلات
            if st.button("🗑️ حذف كافة السجلات نهائياً", key="clear_all_logs"):
                try:
                    c.execute("DELETE FROM logs")
                    conn.commit()
                    st.success("✅ تم مسح سجل العمليات بالكامل")
                    time.sleep(0.5)
                    st.rerun()
                except Exception as e:
                    st.error(f"خطأ أثناء الحذف: {e}")
        st.divider()
        # ✅ عرض السجلات مع تفاصيل التدقيق
                # عرض السجلات (كود آمن يمنع الخطأ)
        try:
            df_l = pd.read_sql("""
                SELECT user as 'الموظف', 
                       action as 'الإجراء', 
                       details as 'التفاصيل', 
                       CASE 
                           WHEN old_value IS NOT NULL AND new_value IS NOT NULL 
                           THEN old_value || ' → ' || new_value 
                           ELSE details 
                       END as 'التغيير',
                       timestamp as 'الوقت' 
                FROM logs ORDER BY id DESC LIMIT 100
            """, conn)
        except:
            # في حال عدم وجود الأعمدة الجديدة، اعرض البيانات القديمة فقط
            df_l = pd.read_sql("""
                SELECT user as 'الموظف', 
                       action as 'الإجراء', 
                       details as 'التفاصيل', 
                       timestamp as 'الوقت' 
                FROM logs ORDER BY id DESC LIMIT 100
            """, conn)
        if not df_l.empty:
            st.dataframe(df_l, use_container_width=True)
        else:
            st.info("سجل العمليات فارغ حالياً.")

           # ==================== تبويب الاستعلامات الذكية ====================
    with tab_inquiry:
        # 🟢 CSS للمحاذاة من اليمين لليسار
        st.markdown("""
        <style>
        .stApp [data-testid="stVerticalBlock"] > div:first-child { direction: rtl !important; text-align: right !important; }
        .stMarkdown h1, .stMarkdown h2, .stMarkdown h3, .stMarkdown h4, .stMarkdown p, .stMarkdown label { direction: rtl !important; text-align: right !important; }
        input[type="text"], input[type="search"], select, textarea, .stTextInput > div > div > input, .stSelectbox > div > div > select { direction: rtl !important; text-align: right !important; }
        .stButton > button { direction: rtl !important; }
        [data-testid="stMetric"] { direction: rtl !important; text-align: center !important; }
        [data-testid="stMetricLabel"] { direction: rtl !important; }
        [data-testid="stMetricValue"] { direction: ltr !important; text-align: center !important; }
        table, th, td, [data-testid="stDataFrame"] { direction: rtl !important; text-align: right !important; }
        .stColumns > div { direction: rtl !important; }
        </style>
        """, unsafe_allow_html=True)

        st.markdown('<h2 style="text-align: right; direction: rtl;">🔎 نظام الاستعلامات الذكية والتحليلات</h2>', unsafe_allow_html=True)
        
        # 🔑 تهيئة الحالات في الجلسة (تمت إضافة q_city)
        if 'query_df' not in st.session_state: st.session_state.query_df = None
        if 'run_query' not in st.session_state: st.session_state.run_query = False
        for k in ['q_role', 'q_pref', 'q_abl', 'q_assigned', 'q_city', 'q_search']:
            if k not in st.session_state: st.session_state[k] = "الكل" if k != 'q_search' else ""

        # 1️⃣ مصدر البيانات
        st.markdown('<h3 style="text-align: right; direction: rtl;">📂 مصدر البيانات</h3>', unsafe_allow_html=True)
        source_options = ["جميع المعلمين (حسب الوضع الحالي)"]
        if st.session_state.system_mode == "tawjihi": source_options.insert(0, "نظام الثانوية العامة")
        elif st.session_state.system_mode == "tawzif": source_options.insert(0, "نظام التوظيف")
        st.selectbox("", source_options, index=0, label_visibility="collapsed", disabled=True)

        # 2️⃣ تقارير سريعة
        st.markdown('<h3 style="text-align: right; direction: rtl;">⚡ تقارير سريعة</h3>', unsafe_allow_html=True)
        q1, q2, q3, q4, q5, q6 = st.columns(6)
        def run_quick_report(role, pref, abl, assign, city="الكل"):
            st.session_state.update({'q_role': role, 'q_pref': pref, 'q_abl': abl, 'q_assigned': assign, 'q_city': city})
            st.session_state.run_query = True
            st.rerun()

        with q1:
            if st.button("👔 مدراء راغبين ويصلحون", use_container_width=True): run_quick_report("مدير مدرسة", "يرغب", "يصلح", "الكل")
        with q2:
            if st.button("📋 سكرتارية راغبين ويصلحون", use_container_width=True): run_quick_report("سكرتير", "يرغب", "يصلح", "الكل")
        with q3:
            if st.button("👨‍🏫 معلمون راغبون ويصلحون", use_container_width=True): run_quick_report("معلم", "يرغب", "يصلح", "الكل")
        with q4:
            if st.button("🔑 آذن راغبون ويصلحون", use_container_width=True): run_quick_report("آذن", "يرغب", "يصلح", "الكل")
        with q5:
            if st.button("✅ جميع المكلفين حالياً", use_container_width=True): run_quick_report("الكل", "الكل", "الكل", "مكلف")
        with q6:
            if st.button("⏳ غير مكلفين (متاحين)", use_container_width=True): run_quick_report("الكل", "الكل", "الكل", "غير مكلف")

        st.divider()

        # 3️⃣ الفلاتر المخصصة (تمت إضافة مكان السكن)
        st.markdown('<h3 style="text-align: right; direction: rtl;">⚙️ فلترة متقدمة</h3>', unsafe_allow_html=True)
        
        # جلب قائمة المدن/السكن لملء القائمة المنسدلة ديناميكياً
        df_cache = get_cached_teachers()
        cities_opts = ["الكل"] + sorted(df_cache['city'].dropna().unique().tolist()) if not df_cache.empty and 'city' in df_cache.columns else ["الكل"]
        
        f1, f2, f3, f4, f5 = st.columns(5)
        with f1:
            st.session_state.q_role = st.selectbox("🎯 الوظيفة/المهمة:", ["الكل", "مدير مدرسة", "سكرتير", "معلم", "آذن", "غير محدد"], index=0 if st.session_state.q_role=="الكل" else ["الكل","مدير مدرسة","سكرتير","معلم","آذن","غير محدد"].index(st.session_state.q_role), key="sel_role")
        with f2:
            st.session_state.q_pref = st.selectbox("💭 الرغبة:", ["الكل", "يرغب", "لا يرغب", "غير محدد"], index=0 if st.session_state.q_pref=="الكل" else ["الكل","يرغب","لا يرغب","غير محدد"].index(st.session_state.q_pref), key="sel_pref")
        with f3:
            st.session_state.q_abl = st.selectbox("🛡️ صلاحية المراقبة:", ["الكل", "يصلح", "لا يصلح", "غير محدد"], index=0 if st.session_state.q_abl=="الكل" else ["الكل","يصلح","لا يصلح","غير محدد"].index(st.session_state.q_abl), key="sel_abl")
        with f4:
            st.session_state.q_assigned = st.selectbox("📌 حالة التكليف:", ["الكل", "مكلف", "غير مكلف"], index=0 if st.session_state.q_assigned=="الكل" else ["الكل","مكلف","غير مكلف"].index(st.session_state.q_assigned), key="sel_assign")
        with f5:
            st.session_state.q_city = st.selectbox("🏡 مكان السكن:", cities_opts, index=0 if st.session_state.q_city=="الكل" else cities_opts.index(st.session_state.q_city), key="sel_city")

        st.session_state.q_search = st.text_input("🔍 بحث حر (اسم، هوية، مدرسة، قاعة، جوال):", value=st.session_state.q_search, key="txt_search")

        if st.button("🚀 تنفيذ الاستعلام", type="primary", use_container_width=True):
            st.session_state.run_query = True

        # 🔍🔍 تنفيذ الاستعلام وحفظه في الجلسة
        if st.session_state.run_query:
            df = get_cached_teachers()
            if df.empty:
                st.warning("⚠️ قاعدة البيانات فارغة أو لم يتم تحميل البيانات بعد.")
                st.session_state.query_df = pd.DataFrame()
            else:
                mask = pd.Series([True] * len(df))
                if st.session_state.q_role != "الكل": mask &= (df['current_job'] == st.session_state.q_role)
                if st.session_state.q_pref != "الكل": mask &= (df['preference'] == st.session_state.q_pref)
                if st.session_state.q_abl != "الكل": mask &= (df['ability'] == st.session_state.q_abl)
                if st.session_state.q_city != "الكل": mask &= (df['city'] == st.session_state.q_city) # 🟢 فلتر السكن الجديد
                
                hall_valid = (df['hall'].astype(str).str.len() > 0) & (df['hall'] != "nan") & (df['hall'].notna())
                if st.session_state.q_assigned == "مكلف": mask &= hall_valid
                elif st.session_state.q_assigned == "غير مكلف": mask &= ~hall_valid

                if st.session_state.q_search:
                    mask &= df.astype(str).apply(lambda col: col.str.contains(st.session_state.q_search, case=False, na=False)).any(axis=1)

                st.session_state.query_df = df[mask].copy()
            st.session_state.run_query = False

        # 📊 عرض النتائج والتصدير
if st.session_state.query_df is not None and not st.session_state.query_df.empty:
    df_res = st.session_state.query_df
    total = len(df_res)
    hall_valid_mask = (df_res['hall'].astype(str).str.len() > 0) & (df_res['hall'] != "nan") & (df_res['hall'].notna())
    assigned = len(df_res[hall_valid_mask])
    unassigned = total - assigned
    pct = (assigned / max(total, 1)) * 100

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("📊 إجمالي النتائج", total)
    c2.metric("✅ المكلفون", assigned)
    c3.metric("⏳ غير مكلفين", unassigned)
    c4.metric("📈 نسبة الإنجاز", f"{pct:.1f}%")

    st.markdown('<h4 style="text-align: right; direction: rtl;">📈 توزيع الرغبة بين النتائج</h4>', unsafe_allow_html=True)
    st.bar_chart(df_res['preference'].value_counts(), horizontal=True)

    st.markdown('<h4 style="text-align: right; direction: rtl;">📋 جدول النتائج التفصيلي</h4>', unsafe_allow_html=True)
    
    # ✅✅✅ الجديد: إنشاء نسخة للعرض وتغيير أسماء الأعمدة للعربي
    df_display = df_res.copy()
    arabic_map = {
        'name': 'الاسم',
        'id': 'رقم الهوية',
        'current_job': 'الوظيفة',
        'preference': 'الرغبة',
        'ability': 'الصلاحية',
        'hall': 'القاعة',
        'city': 'السكن',
        'phone': 'الجوال',
        'school': 'المدرسة',
        'role': 'المهمة',
        'hall_city': 'مدينة القاعة',
        'notes': 'ملاحظات'
    }
    # تطبيق التغيير فقط على الأعمدة الموجودة فعلياً
    df_display = df_display.rename(columns={k: v for k, v in arabic_map.items() if k in df_display.columns})

    # ترتيب الأعمدة اللي بدك تظهر في الجدول
    display_cols = ['الاسم', 'رقم الهوية', 'الوظيفة', 'الرغبة', 'الصلاحية', 'القاعة', 'السكن', 'الجوال']
    safe_cols = [c for c in display_cols if c in df_display.columns]
    
    st.dataframe(df_display[safe_cols], use_container_width=True)

    st.divider()
    if st.button("📥 تصدير التقرير إلى Excel", type="secondary"):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            # ✅ نستخدم df_display عشان يطلع بالإكسل بنفس الأسماء العربية
            df_export = df_display[safe_cols].copy()
            df_export.to_excel(writer, index=False, sheet_name='نتائج_الاستعلام')
            
            wb = writer.book
            ws = writer.sheets['نتائج_الاستعلام']
            header_fmt = wb.add_format({'bold': True, 'bg_color': '#1a1c23', 'font_color': '#00ffcc', 'border': 1, 'align': 'center'})
            cell_fmt = wb.add_format({'border': 1, 'align': 'right'})
            ws.right_to_left()
            ws.set_landscape()
            ws.fit_to_pages(1, 0)
            for col_num, col_name in enumerate(df_export.columns):
                ws.write(0, col_num, col_name, header_fmt)
                max_len = max(df_export[col_name].astype(str).str.len().max(), len(col_name)) + 5
                ws.set_column(col_num, col_num, min(max_len, 40), cell_fmt)
        
        st.download_button(
            label="📥 تحميل ملف Excel",
            data=output.getvalue(),
            file_name=f"تقرير_استعلام_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="dl_excel_inquiry_v2"
        )
elif st.session_state.query_df is not None:
    st.info("لا توجد نتائج مطابقة للشروط المختارة.")

# ============================================================================
# ✨✨✨ نظام تصحيح الثانوية العامة - وحدة مستقلة تماماً ✨✨✨
# ============================================================================
# 🔴 هذا القسم يظهر فقط إذا كان الوضع == "tasheeh"

if st.session_state.get('system_mode') == "tasheeh":
    
        # 1️⃣ إنشاء جداول التخزين الدائم في قاعدة البيانات
    c.execute('''CREATE TABLE IF NOT EXISTS tasheeh_teachers (
        id TEXT PRIMARY KEY, name TEXT, subject TEXT, city TEXT, 
        school TEXT, phone TEXT, relative TEXT
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS tasheeh_halls (
        hall_name TEXT PRIMARY KEY, city TEXT
    )''')
    
    # ✅ إضافة جدول التكليفات المفقود (هذا هو الحل)
    c.execute('''CREATE TABLE IF NOT EXISTS tasheeh_assignments (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        teacher_id TEXT, teacher_name TEXT, subject TEXT,
        hall_name TEXT, hall_city TEXT, exam_name TEXT,
        exam_date TEXT, exam_day TEXT, 
        school TEXT, city TEXT,
        created_at TEXT, created_by TEXT
    )''')
    
    conn.commit()
    
    # 2️⃣ تحميل البيانات تلقائياً من القاعدة عند فتح الموقع
    if 'tasheeh_teachers' not in st.session_state:
        try:
            st.session_state['tasheeh_teachers'] = pd.read_sql("SELECT * FROM tasheeh_teachers", conn)
        except: st.session_state['tasheeh_teachers'] = pd.DataFrame()
    if 'tasheeh_halls' not in st.session_state:
        try:
            st.session_state['tasheeh_halls'] = pd.read_sql("SELECT * FROM tasheeh_halls", conn)
        except: st.session_state['tasheeh_halls'] = pd.DataFrame()

    # 3️⃣ دالة المزامنة الذكية (تحديث + إضافة بدون تكرار)
    def sync_tasheeh_data():
        try:
            with st.spinner("🔄 جاري المزامنة والتنظيف العميق..."):
                df_t = pd.read_csv(TEACHERS_URL, dtype=str)
                df_t.columns = df_t.columns.str.strip().str.lower()
                rename_map = {
                    'رقم الهوية': 'id', 'الاسم': 'name', 'المبحث': 'subject',
                    'مكان سكن المعلم': 'city', 'اسم المدرسة': 'school', 
                    'رقم جواله': 'phone', 'هل له قريب مباشر او لا': 'relative'
                }
                df_t = df_t.rename(columns={k:v for k,v in rename_map.items() if k in df_t.columns})

                # 🟢 1. تنظيف رقم الهوية من أي مسافات خفية أو زائدة (مهم جداً)
                df_t['id'] = df_t['id'].astype(str).str.strip()
                df_t['id'] = df_t['id'].str.replace(' ', '') # إزالة المسافات من داخل الرقم أيضاً
                
                # 🟢 2. حذف التكرار في ملف الإكسل نفسه (نحتفظ بالصف الأول/المدرسة الأولى فقط)
                before_csv = len(df_t)
                df_t = df_t.drop_duplicates(subset=['id'], keep='first')
                after_csv = len(df_t)
                if before_csv > after_csv:
                    st.warning(f"⚠️ تم تجاهل `{before_csv - after_csv}` تكرار موجود في ملف الإكسل")

                for _, r in df_t.iterrows():
                    tid = str(r.get('id','')).strip().replace(' ', '')
                    if not tid: continue
                    c.execute("""INSERT OR REPLACE INTO tasheeh_teachers 
                                 (id, name, subject, city, school, phone, relative) 
                                 VALUES (?,?,?,?,?,?,?)""",
                              (tid, str(r.get('name','')), str(r.get('subject','')),
                               str(r.get('city','')), str(r.get('school','')), 
                               str(r.get('phone','')), str(r.get('relative',''))))
                conn.commit()

                # 🟢 3. تنظيف قاعدة البيانات من الداخل (إزالة المسافات من الأرقام المخزنة سابقاً)
                c.execute("UPDATE tasheeh_teachers SET id = TRIM(id)")
                
                # 🟢 4. حذف المكررات المتبقية في قاعدة البيانات (نحتفظ بأول ظهور فقط)
                c.execute("""
                    DELETE FROM tasheeh_teachers
                    WHERE rowid NOT IN (
                        SELECT MIN(rowid)
                        FROM tasheeh_teachers
                        GROUP BY id
                    )
                """)
                conn.commit()

                df_h = pd.read_csv(HALLS_URL, dtype=str)
                df_h.columns = df_h.columns.str.strip().str.upper()
                for _, r in df_h.iterrows():
                    hname = str(r.get('ZHALL','')).strip()
                    if not hname: continue
                    c.execute("INSERT OR REPLACE INTO tasheeh_halls (hall_name, city) VALUES (?,?)",
                              (hname, str(r.get('ZLOC',''))))
                conn.commit()

                st.session_state['tasheeh_teachers'] = pd.read_sql("SELECT * FROM tasheeh_teachers", conn)
                st.session_state['tasheeh_halls'] = pd.read_sql("SELECT * FROM tasheeh_halls", conn)
            st.success(f"✅ تم التحديث بنجاح! (العدد النهائي النظيف: `{len(st.session_state['tasheeh_teachers'])}`)")
            st.rerun()
        except Exception as e:
            st.error(f"❌ خطأ أثناء المزامنة: {e}")

    def generate_tasheeh_letter(data, exam_name):
        if not os.path.exists(TEMPLATE_NAME):
            return None
        doc = Document(TEMPLATE_NAME)
        
        # ✅ إصلاح 1: ZWORK يأخذ قيمة school بشكل صحيح مع fallback
        repls = {
            'ZNAME': str(data.get('name', '---')),
            'ZID': str(data.get('id', '---')),
            'ZTEST': str(exam_name),
            'ZHALL': str(data.get('hall_name', '---')),
            'ZLOC': str(data.get('hall_city', '---')),
            'ZWORK': str(data.get('school', data.get('ZWORK', '---'))),  # ✅ fallback آمن
            'ZSCHOOL': str(data.get('school', '---')),  # ✅ حقل إضافي للقالب
            'ZCITY': str(data.get('city', '---')),
            'ZSUBJECT': str(data.get('subject', '---')),
            'ZDATE': str(data.get('exam_date', '---'))
        }
        
        # ✅ دالة مساعدة للاستبدال في الفقرات والجداول
        def replace_in_element(element, k, v):
            for run in element.runs:
                if k in run.text:
                    run.text = run.text.replace(k, str(v))
                    run.bold = True
        
        for p in doc.paragraphs:
            for k, v in repls.items():
                if k in p.text:
                    replace_in_element(p, k, v)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for k, v in repls.items():
                            if k in p.text:
                                replace_in_element(p, k, v)
        
        return doc
    
    st.markdown("""
        <div style="background: linear-gradient(135deg, #1a1c23 0%, #2d3748 100%);
                    padding: 20px; border-radius: 15px; border: 2px solid #00ffcc;
                    margin: 20px 0; text-align: center;">
            <h2 style="color: #00ffcc; margin: 0;">✨ نظام تصحيح الثانوية العامة ✨</h2>
            <p style="color: #bbb; margin: 10px 0 0 0;">توزيع المصححين حسب المبحث والقاعة</p>
        </div>
        
        <style>
        /* RTL - اجبار كل العناصر على الاتجاه من اليمين لليسار */
        * {
            direction: rtl !important;
        }
        
        body, .stApp, main, section, div[class*="st-"] {
            direction: rtl !important;
            text-align: right !important;
        }
        
        /* النصوص والعناوين */
        h1, h2, h3, h4, h5, h6, p, span, label {
            direction: rtl !important;
            text-align: right !important;
        }
        
        /* حقول الإدخال */
        input[type="text"], input[type="number"], input[type="date"], 
        textarea, select, [data-baseweb="input"], [data-baseweb="select"] {
            direction: rtl !important;
            text-align: right !important;
        }
        
        /* الأزرار */
        button, [role="button"] {
            direction: rtl !important;
        }
        
        /* الجداول */
        table, thead, tbody, tr, th, td {
            direction: rtl !important;
            text-align: right !important;
        }
        
        /* التبويبات */
        [data-testid="stTabs"], [data-testid="stTab"], .stTabs {
            direction: rtl !important;
        }
        
        /* الصناديق والمحتوى */
        .stAlert, .stInfo, .stSuccess, .stWarning, .stError,
        [data-testid="stAlert"], [data-testid="stMarkdown"] {
            direction: rtl !important;
            text-align: right !important;
        }
        
        /* القوائم المنسدلة */
        [data-testid="stSelectbox"], [data-testid="stMultiselect"] {
            direction: rtl !important;
        }
        </style>
    """, unsafe_allow_html=True)
           
    
    corr_tab1, corr_tab2, corr_tab3, corr_tab4 = st.tabs([
        "📥 رفع البيانات", "🔄 التوزيع التلقائي", "📄 كتب التكليف", "📜 سجل العمليات"
    ])
    
    # ==================== تبويب 1: رفع البيانات والمزامنة ====================
    with corr_tab1:
        st.markdown("### 📥 إدارة البيانات وقالب التكليف")
        
        # ✅ تم حذف رفع القالب لأنه أصبح على جيت هب
       
        
        st.divider()
        st.markdown("**2️⃣ المزامنة مع Google Sheets**")
        st.info("💡 البيانات محفوظة تلقائياً في النظام. اضغط هنا فقط إذا أضفت/عدلت بيانات في ملف الإكسل الخارجي.")
        
        if st.button("🔄 مزامنة وتحديث البيانات من Google Sheets", type="primary", use_container_width=True):
            sync_tasheeh_data()
            
        if not st.session_state['tasheeh_teachers'].empty:
            st.markdown(f"📊 **عدد المعلمين المخزنين حالياً:** `{len(st.session_state['tasheeh_teachers'])}`")
            st.dataframe(st.session_state['tasheeh_teachers'].head(), use_container_width=True)
            
        st.divider()
        st.markdown("### 🔍 فحص وحذف التكرار (فحص مباشر من قاعدة البيانات)")
        
        # جلب البيانات مباشرة من قاعدة البيانات
        df_raw = pd.read_sql("SELECT rowid, id, name, subject FROM tasheeh_teachers", conn)
        
        if not df_raw.empty:
            # تنظيف الأرقام وإزالة المسافات
            df_raw['id_trimmed'] = df_raw['id'].astype(str).str.strip()
            
            # البحث عن المكررات
            duplicates = df_raw[df_raw.duplicated(subset=['id_trimmed'], keep=False)]
            
            if len(duplicates) > 0:
                st.error(f"🚨 تم العثور على `{len(duplicates)}` سجل مكرر!")
                st.write("**تفاصيل المكررات:**")
                st.dataframe(duplicates[['rowid', 'id', 'id_trimmed', 'name', 'subject']], use_container_width=True)
                
                # عرض الإحصائيات
                total = len(df_raw)
                unique = df_raw['id_trimmed'].nunique()
                st.warning(f"📊 الإجمالي: `{total}` | الفريد: `{unique}` | المكرر: `{total - unique}`")
            else:
                st.success("✅ لا يوجد تكرار في قاعدة البيانات")
                st.write(f"📊 إجمالي المعلمين: `{len(df_raw)}`")
        else:
            st.info("⚠️ قاعدة البيانات فارغة")

        if st.button("🗑️ حذف المكررات الآن", type="secondary", use_container_width=True):
            try:
                # عد السجلات قبل الحذف
                before = pd.read_sql("SELECT COUNT(*) FROM tasheeh_teachers", conn).iloc[0, 0]
                
                # حذف المكررات (الاحتفاظ بأول rowid فقط)
                c.execute("""
                    DELETE FROM tasheeh_teachers
                    WHERE rowid NOT IN (
                        SELECT MIN(rowid)
                        FROM tasheeh_teachers
                        GROUP BY TRIM(id)
                    )
                """)
                conn.commit()
                
                # عد السجلات بعد الحذف
                after = pd.read_sql("SELECT COUNT(*) FROM tasheeh_teachers", conn).iloc[0, 0]
                deleted = before - after
                
                # تحديث الذاكرة فوراً
                st.session_state['tasheeh_teachers'] = pd.read_sql("SELECT * FROM tasheeh_teachers", conn)
                st.cache_data.clear()
                
                if deleted > 0:
                    st.success(f"✅✅✅ تم حذف `{deleted}` سجل مكرر بنجاح!")
                    st.balloons()
                    st.rerun()
                else:
                    st.warning("⚠️ لم يتم العثور على مكررات للحذف")
                    
            except Exception as e:
                st.error(f"❌ خطأ: {e}")
    # ==================== تبويب 2: التوزيع التلقائي ====================
    with corr_tab2:
        if st.session_state['tasheeh_teachers'].empty or st.session_state['tasheeh_halls'].empty:
            st.warning("⚠️ يرجى مزامنة البيانات أولاً من تبويب 'رفع البيانات'")
        else:
            teachers = st.session_state['tasheeh_teachers']
            halls = st.session_state['tasheeh_halls']
            
            st.markdown("### ⚙️ إعدادات التوزيع")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                # 1️⃣ اختيار القاعة من بيانات الإكسل
                hall_opts = sorted(halls['hall_name'].unique().tolist())
                selected_hall = st.selectbox("🏫 اختر القاعة:", hall_opts, index=0)
                
            with col2:
                # 2️⃣ اختيار المادة من بيانات المعلمين
                subj_opts = sorted(teachers['subject'].dropna().unique().tolist())
                selected_subj = st.selectbox("📚 اختر المادة:", subj_opts, index=0)
                
            with col3:
                # 3️⃣ اختيار التاريخ يدوياً
                default_date = datetime(2026, 6, 20)
                selected_date = st.date_input("📅 تاريخ بداية التصحيح:", value=default_date, min_value=datetime(2026, 1, 1))
            
            # تنسيق التاريخ واليوم
            date_str = selected_date.strftime("%Y/%m/%d")
            days_ar = ["الإثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت", "الأحد"]
            day_name = days_ar[selected_date.weekday()]
            
            # عرض ملخص سريع قبل التوزيع
            hall_row = halls[halls['hall_name'] == selected_hall]
            hall_city = hall_row['city'].values[0] if not hall_row.empty else "غير محدد"
            st.info(f"📍 القاعة: `{selected_hall}` | 🌆 المدينة: `{hall_city}` | 📅 التاريخ: `{date_str} ({day_name})`")

            if st.button("🚀 توزيع المصححين", type="primary", use_container_width=True):
                # تصفية المعلمين (استبعاد من له قريب + تصفية حسب المادة)
                teachers_filtered = teachers[teachers['relative'].astype(str).str.lower() != 'true']
                pool = teachers_filtered[teachers_filtered['subject'] == selected_subj]
                
                if pool.empty:
                    st.warning(f"⚠️ لا يوجد معلمين متاحين للمادة: `{selected_subj}`")
                else:
                    assignments = []
                    for _, t in pool.iterrows():
                        assignments.append({
                            'id': t.get('id',''), 
                            'name': t.get('name',''), 
                            'subject': t.get('subject',''),
                            'hall_name': selected_hall,      # القاعة المختارة
                            'hall_city': hall_city,          # المدينة المجلوبة تلقائياً من الإكسل
                            'exam_name': selected_subj,      # اسم المادة هو اسم الامتحان
                            'exam_date': date_str,           # التاريخ المختار
                            'exam_day': day_name,            # اليوم المحسوب تلقائياً
                            'school': t.get('school',''), 
                            'city': t.get('city',''),
                            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M")
                        })
                    
                    if assignments:
                        st.session_state['tasheeh_assignments'] = assignments
                        add_log("توزيع تصحيح", f"توزيع {len(assignments)} معلم - {selected_subj} في {selected_hall} بتاريخ {date_str}")
                        st.success(f"✅ تم توزيع {len(assignments)} معلم بنجاح!")
                        
                        for a in assignments:
                            try:
                                c.execute("""INSERT INTO tasheeh_assignments 
                                             (teacher_id,teacher_name,subject,hall_name,hall_city,exam_name,
                                              exam_date,exam_day,created_at,created_by) 
                                             VALUES (?,?,?,?,?,?,?,?,?,?)""",
                                         (a['id'],a['name'],a['subject'],a['hall_name'],a['hall_city'],a['exam_name'],
                                          a['exam_date'],a['exam_day'],a['timestamp'],st.session_state.username))
                            except: pass
                        conn.commit()
                    else:
                        st.error("❌ لم يتم التوزيع.")
    
    # ==================== تبويب 3: كتب التكليف ====================
    with corr_tab3:
        
        # 1. تحميل البيانات الأساسية
        assigns = st.session_state.get('tasheeh_assignments', [])
        
        if not assigns:
            st.info("📌 لم يتم توزيع أي تكليفات بعد. اذهب لتبويب التوزيع التلقائي.")
        else:
            st.markdown("### 📊 إحصائيات التكليفات")
            
            # 2. فلترة المواد
            teachers_df = st.session_state['tasheeh_teachers']
            subjects_list = sorted(teachers_df['subject'].dropna().unique().tolist()) if not teachers_df.empty else []
            
            col_f1, col_f2 = st.columns([1, 2])
            with col_f1:
                filter_subj = st.selectbox("🔍 عرض إحصائيات مادة محددة:", ["الكل"] + subjects_list, index=0, key="filter_subj_selectbox")
            
            df_assigns = pd.DataFrame(assigns)
            
            if filter_subj != "الكل":
                pool_count = len(teachers_df[teachers_df['subject'] == filter_subj])
                assigned_df = df_assigns[df_assigns['subject'] == filter_subj] if not df_assigns.empty else pd.DataFrame()
            else:
                pool_count = len(teachers_df)
                assigned_df = df_assigns
            
            assigned_count = len(assigned_df)
            remaining_count = pool_count - assigned_count

            # 3. عرض المقاييس
            c_m1, c_m2, c_m3 = st.columns(3)
            with c_m1: st.metric("📚 إجمالي المعلمين (المادة)", pool_count)
            with c_m2: st.metric("✅ تم تكليفهم", assigned_count)
            with c_m3: st.metric("⏳ المتبقي للتكليف", remaining_count)

            st.divider()

            # 4. عرض الجدول
            # 4. عرض الجدول
            st.markdown(f"### 📋 القائمة الحالية: {filter_subj}")
            if not assigned_df.empty:
                # ✅ تحويل أسماء الأعمدة للعربي
                df_display = assigned_df.copy()
                arabic_cols = {
                    'name': 'الاسم',
                    'subject': 'المبحث',
                    'hall_name': 'القاعة',
                    'hall_city': 'المدينة',
                    'exam_date': 'التاريخ',
                    'id': 'رقم الهوية',
                    'exam_day': 'اليوم'
                }
                df_display = df_display.rename(columns={k: v for k, v in arabic_cols.items() if k in df_display.columns})
                
                display_cols = ['الاسم', 'المبحث', 'القاعة', 'المدينة', 'التاريخ']
                safe_cols = [c for c in display_cols if c in df_display.columns]
                st.dataframe(df_display[safe_cols], use_container_width=True)
            else:
                st.info("لا يوجد تكليفات لهذه المادة.")

            st.divider()

            # 5. أزرار التحكم (تحميل وحذف وتصدير)
            st.markdown("### ⚙️ إدارة وتصدير")
            
            # ✅ تعريف الأعمدة مرة واحدة فقط
            col_btn1, col_btn2 = st.columns(2)
            
            # === العمود 1: تحميل الوورد ===
            with col_btn1:
                btn_key_word = f"btn_word_export_{filter_subj}"
                if st.button("📥 تحميل وورد للمادة الحالية", type="primary", use_container_width=True, disabled=assigned_df.empty, key=btn_key_word):
                    if not os.path.exists(TEMPLATE_NAME):
                        st.error("❌ القالب غير موجود")
                    else:
                        with st.spinner("جاري الإنشاء..."):
                            try:
                                final_doc = Document(TEMPLATE_NAME)
                                final_doc._body.clear_content()
                                current_list = assigned_df.to_dict('records') 
                                
                                for i, a in enumerate(current_list):
                                    temp_doc = Document(TEMPLATE_NAME)
                                    # ✅ إصلاح 1: ZWORK يأخذ school بشكل صحيح
                                    repls = {
                                        'ZNAME': str(a.get('name', '---')), 
                                        'ZID': str(a.get('id', '---')),
                                        'ZTEST': str(a.get('exam_name', '---')), 
                                        'ZHALL': str(a.get('hall_name', '---')),
                                        'ZLOC': str(a.get('hall_city', '---')), 
                                        'ZWORK': str(a.get('school', a.get('ZWORK', '---'))),  # ✅ fallback آمن
                                        'ZSCHOOL': str(a.get('school', '---')),  # ✅ حقل إضافي
                                        'ZCITY': str(a.get('city', '---')),
                                        'ZSUBJECT': str(a.get('subject', '---')),
                                        'ZDATE': str(a.get('exam_date', '---')) 
                                    }
                                    for p in temp_doc.paragraphs:
                                        for k, v in repls.items():
                                            if k in p.text:
                                                for run in p.runs:
                                                    if k in run.text:
                                                        run.text = run.text.replace(k, v)
                                                        run.bold = True
                                    for table in temp_doc.tables:
                                        for row in table.rows:
                                            for cell in row.cells:
                                                for p in cell.paragraphs:
                                                    for k, v in repls.items():
                                                        if k in p.text:
                                                            for run in p.runs:
                                                                if k in run.text:
                                                                    run.text = run.text.replace(k, v)
                                                                    run.bold = True
                                    
                                    elements = [el for el in temp_doc.element.body if not el.tag.endswith('sectPr')]
                                    for element in elements:
                                        final_doc.element.body.append(copy.deepcopy(element))
                                    
                                    if i < len(current_list) - 1:
                                        p = OxmlElement('w:p')
                                        r = OxmlElement('w:r')
                                        br = OxmlElement('w:br')
                                        br.set(qn('w:type'), 'page')
                                        r.append(br)
                                        p.append(r)
                                        final_doc.element.body.append(p)
                                
                                out = io.BytesIO()
                                final_doc.save(out)
                                out.seek(0)
                                st.success(f"✅ تم إنشاء ملف {filter_subj} بنجاح")
                                st.download_button(
                                    label="📥 تحميل الآن",
                                    data=out.getvalue(),
                                    file_name=f"تكليفات_{filter_subj}_{datetime.now().strftime('%Y%m%d')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"dl_word_{filter_subj}"
                                )
                            except Exception as e:
                                st.error(f"خطأ: {e}")

            # === العمود 2: الحذف وتصدير الإكسل ===
            with col_btn2:
                st.markdown("#### 📊 خيارات التصدير")
                
                # زر الحذف
                del_btn_label = f"🗑️ حذف تكليفات {filter_subj}" if filter_subj != "الكل" else "🗑️ حذف جميع التكليفات"
                del_btn_key = f"btn_delete_{filter_subj}"
                
                if st.button(del_btn_label, type="secondary", use_container_width=True, disabled=assigned_df.empty, key=del_btn_key):
                    if filter_subj == "الكل":
                        c.execute("DELETE FROM tasheeh_assignments")
                    else:
                        c.execute("DELETE FROM tasheeh_assignments WHERE subject=?", (filter_subj,))
                    conn.commit()
                    
                    if filter_subj == "الكل":
                        st.session_state['tasheeh_assignments'] = []
                    else:
                        st.session_state['tasheeh_assignments'] = [a for a in st.session_state['tasheeh_assignments'] if a['subject'] != filter_subj]
                    
                    st.success(f"✅ تم حذف تكليفات {filter_subj}")
                    st.rerun()
                
                st.divider()
                
                # ✅ زر الإكسل المنسق
                excel_btn_key = f"btn_excel_export_{filter_subj}"
                
                if st.button("📥 تصدير إكسل منسق", type="primary", use_container_width=True, disabled=assigned_df.empty, key=excel_btn_key):
                    if assigned_df.empty:
                        st.warning("⚠️ لا توجد بيانات لتصديرها!")
                    else:
                        arabic_map = {
                            'id': 'رقم الهوية', 'name': 'اسم المصحح', 'subject': 'المبحث',
                            'hall_name': 'القاعة', 'hall_city': 'المدينة', 'exam_name': 'الامتحان',
                            'exam_date': 'التاريخ', 'exam_day': 'اليوم', 'city': 'مكان السكن', 
                            'school': 'المدرسة', 'timestamp': 'وقت التكليف'
                        }
                        valid_cols = [c for c in assigned_df.columns if c in arabic_map]
                        df_export = assigned_df[valid_cols].rename(columns=arabic_map)

                        out = io.BytesIO()
                        with pd.ExcelWriter(out, engine='xlsxwriter') as writer:
                            df_export.to_excel(writer, index=False, sheet_name='تكليفات التصحيح')
                            workbook = writer.book
                            worksheet = writer.sheets['تكليفات التصحيح']

                            header_fmt = workbook.add_format({'font_size': 14, 'bold': True, 'align': 'center', 'valign': 'vcenter', 'bg_color': '#1a1c23', 'font_color': '#00ffcc', 'border': 1, 'text_wrap': True})
                            cell_fmt = workbook.add_format({'font_size': 14, 'bold': True, 'align': 'right', 'valign': 'vcenter', 'border': 1, 'text_wrap': True})

                            worksheet.right_to_left()
                            worksheet.set_landscape()
                            worksheet.fit_to_pages(1, 0)
                            worksheet.set_default_row(height=28)

                            for col_num, value in enumerate(df_export.columns):
                                worksheet.write(0, col_num, value, header_fmt)
                            for row_num in range(len(df_export)):
                                for col_num in range(len(df_export.columns)):
                                    worksheet.write(row_num + 1, col_num, df_export.iloc[row_num, col_num], cell_fmt)
                            for idx, col in enumerate(df_export.columns):
                                max_len = max(df_export[col].astype(str).map(len).max(), len(str(col))) + 4
                                worksheet.set_column(idx, idx, min(max_len, 35), cell_fmt)

                        out.seek(0)
                        st.download_button(
                            label="📥 تحميل ملف إكسل منسق",
                            data=out.getvalue(),
                            file_name=f"تكليفات_{filter_subj}_{datetime.now().strftime('%Y%m%d')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key=f"dl_excel_{filter_subj}"
                        )
    # ==================== تبويب 4: سجل العمليات ====================
    with corr_tab4:
        st.markdown("### 📜 سجل العمليات الخاص بالتصحيح")
        
        # 🔴🔴 زر حذف السجلات الجديد 🔴🔴
        if is_admin():
            st.warning("⚠️ هذا الزر سيقوم بحذف سجلات التصحيح نهائياً.")
            if st.button("🗑️ حذف سجلات التصحيح نهائياً", type="primary", key="delete_tasheeh_logs_btn"):
                try:
                    # حذف السجلات التي تحتوي على كلمة 'تصحيح' فقط
                    c.execute("DELETE FROM logs WHERE action LIKE '%تصحيح%'")
                    conn.commit()
                    st.success("✅ تم مسح سجلات التصحيح بالكامل")
                    time.sleep(1)
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ خطأ أثناء الحذف: {e}")
        
        st.divider()
        
        # عرض السجلات
        df = pd.read_sql("SELECT user as 'الموظف', action as 'الإجراء', details as 'التفاصيل', timestamp as 'الوقت' FROM logs WHERE action LIKE '%تصحيح%' ORDER BY id DESC", conn)
        if not df.empty:
            st.dataframe(df, use_container_width=True)
        else:
            st.info("لا يوجد سجلات تصحيح حالياً.")
    
    st.stop()


# ============================================================================
# ============================================================================
# 📋 نظام التكليفات الأخرى - وحدة مستقلة تماماً
# ============================================================================

if st.session_state.get('system_mode') == "other_assignments":
    
    # إنشاء قاعدة البيانات والجداول
    conn_other = sqlite3.connect(DB_NAME, check_same_thread=False, timeout=30)
    c_other = conn_other.cursor()
    
    for table_name in ['guards', 'parcels', 'exam_device', 'exam_committee']:
        c_other.execute(f'''CREATE TABLE IF NOT EXISTS {table_name} 
                     (id INTEGER PRIMARY KEY AUTOINCREMENT,
                      zid TEXT, zname TEXT, zjob TEXT, zwork TEXT, 
                      zloc TEXT, zcity TEXT, zdate TEXT, created_at TEXT)''')
        try:
            c_other.execute(f"ALTER TABLE {table_name} ADD COLUMN zdate TEXT")
            conn_other.commit()
        except:
            pass
    conn_other.commit()
        # ✅ إضافة عمود zjob2 للوظيفة الحالية
    for table_name in ['guards', 'parcels', 'exam_device', 'exam_committee']:
        try:
            c_other.execute(f"ALTER TABLE {table_name} ADD COLUMN ZSCHOOL TEXT")
            conn_other.commit()
        except:
            pass  # العمود موجود مسبقاً
    
    def add_other_assignment(table_name, zid, zname, zjob, ZSCHOOL, zwork, zloc, zcity, zdate):
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        c_other.execute(f"""INSERT INTO {table_name} 
                           (zid, zname, zjob, ZSCHOOL, zwork, zloc, zcity, zdate, created_at) 
                           VALUES (?,?,?,?,?,?,?,?,?)""",
                       (zid, zname, zjob, ZSCHOOL, zwork, zloc, zcity, zdate, now))
        conn_other.commit()
        
    def get_other_assignments(table_name):
        return pd.read_sql(f"SELECT * FROM {table_name} ORDER BY id DESC", conn_other)
    
        # ✅ دالة حذف التكليف (مع إصلاح الحذف من قاعدة البيانات)
    def delete_other_assignment(table_name, record_id):
        """حذف تكليف معين"""
        c_other.execute(f"DELETE FROM {table_name} WHERE id=?", (record_id,))
        conn_other.commit()  # ✅ تأكيد الحفظ فوراً
        
    def generate_other_letter(row):
        if not os.path.exists(TEMPLATE_NAME):
            return None
        
        doc = Document(TEMPLATE_NAME)
        
        # ✅ إصلاح 2: إخفاء ZWORK للحرس فقط + إصلاح ZJOBb
        is_guard = str(row.get('zjob', '')).strip() == 'حارس'
        
        repls = {
            'ZID': str(row.get('zid', '---')),
            'ZNAME': str(row.get('zname', '---')),
            'ZJOB': str(row.get('zjob', '---')),         # المهمة (حارس، مرافق...)
            'ZSCHOOL': str(row.get('ZSCHOOL', '---')),  # ✅ الوظيفة الحالية (معلم، إداري...)
            'ZWORK': '' if is_guard else str(row.get('zwork', '---')),       # ✅ إخفاء للحرس
            'ZLOC': str(row.get('zloc', '---')),
            'ZCITY': str(row.get('zcity', '---')),
            'ZDATE': str(row.get('zdate', datetime.now().strftime("%Y/%m/%d")))
        }
        
        for p in doc.paragraphs:
            for k, v in repls.items():
                if k in p.text:
                    for run in p.runs:
                        if k in run.text:
                            run.text = run.text.replace(k, str(v))
                            run.bold = True
        
        for table in doc.tables:
            for row_tbl in table.rows:
                for cell in row_tbl.cells:
                    for p in cell.paragraphs:
                        for k, v in repls.items():
                            if k in p.text:
                                for run in p.runs:
                                    if k in run.text:
                                        run.text = run.text.replace(k, str(v))
                                        run.bold = True
        
        # 🧹 حذف فقرة ZWORK نهائياً إذا كانت فارغة (خاص بالحرس)
        if is_guard:
            for p in doc.paragraphs[:]:  # [:] لتجنب خطأ التعديل أثناء التكرار
                if 'ZWORK' in p.text or (p.text.strip() == ''):
                    p._element.getparent().remove(p._element)
        
        return doc

        # الواجهة الرئيسية
    st.markdown(f"""
        <div style="background: linear-gradient(135deg, #1a1c23 0%, #2d3748 100%); 
                    padding: 20px; border-radius: 15px; border: 2px solid #00ffcc;
                    margin: 20px 0; text-align: center;">
            <h2 style="color: #00ffcc; margin: 0; text-align: center;">📋 نظام التكليفات الأخرى</h2>
            <p style="color: #bbb; margin: 10px 0 0 0; text-align: center;">الحرس | مرافقة الطرود | جهاز الامتحان | لجنة الامتحان</p>
        </div>
        
        <!-- ✅ تنسيق الهيدر العلوي ليكون في المنتصف (فقط في تكليفات أخرى) -->
        <style>
        /* اجبار الاتجاه العام */
        .stApp {{
            direction: rtl !important;
            text-align: right !important;
        }}
        
        /* ✅ الهيدر العلوي في المنتصف فقط لهذا النظام */
        .custom-header-other {{
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            background-color: #1a1c23;
            color: white;
            text-align: center !important;  /* ✅ النص في المنتصف */
            padding: 15px 0;
            z-index: 999999;
            border-bottom: 2px solid #00ffcc;
            line-height: 1.5;
            direction: rtl;
            box-shadow: 0px 4px 10px rgba(0,0,0,0.5);
        }}
        .custom-header-other div {{
            text-align: center !important;  /* ✅ النصوص الداخلية في المنتصف */
            display: block;
            width: 100%;
        }}
        
        /* النصوص والعناوين */
        p, h3, h4, h5, h6, label, span, div {{
            direction: rtl !important;
            text-align: right !important;
        }}
        
        /* حقول الكتابة والأزرار */
        input, textarea, button {{
            direction: rtl !important;
            text-align: right !important;
        }}
        
        /* التبويبات */
        [data-testid="stTabs"] {{
            direction: rtl !important;
            justify-content: flex-end !important;
        }}
        [data-testid="stTab"] {{
            direction: rtl !important;
        }}
        
        /* الجداول */
        [data-testid="stDataFrame"], table, thead, tbody, tr, td, th {{
            direction: rtl !important;
            text-align: right !important;
        }}
        </style>
        
        <!-- ✅ الهيدر الجديد في المنتصف -->
        <div class="custom-header-other">
            <div style="font-weight: bold; font-size: 1.2rem;">إعداد وتصميم : عوض نعمان ريده</div>
            <div style="font-size: 1rem; color: #00ffcc;">قسم الامتحانات - مديرية التربية والتعليم جنوب نابلس</div>
        </div>
    """, unsafe_allow_html=True)
    
    # ✅ تم حذف رفع القالب لأنه أصبح على جيت هب
    
    st.divider()
    
    # التبويبات الأربعة
    tab_guard, tab_parcels, tab_device, tab_committee = st.tabs([
        "🛡️ الحرس", "📦 مرافقة الطرود", "📱 جهاز الامتحان", "👥 لجنة الامتحان"
    ])
    
        # ==================== تبويب الحرس ====================
    with tab_guard:
        st.markdown("### 🛡️ إدارة تكليفات الحرس")
        
        if 'guard_form_clear' not in st.session_state:
            st.session_state.guard_form_clear = False
        
        col1, col2 = st.columns([1, 2])
        with col1:
            st.markdown("**➕ إضافة تكليف جديد**")
            with st.form("add_guard_form"):
                g_zid = st.text_input("رقم الهوية (ZID)", key="g_zid_input")
                g_zname = st.text_input("الاسم (ZNAME)", key="g_zname_input")
                g_zjob = st.text_input("المهمة (ZJOB)", value="حارس", key="g_zjob_input")
                g_ZSCHOOL = st.text_input("الوظيفة الحالية (ZSCHOOL)", value="", key="g_ZSCHOOL_input")  # ✅ key فريد + value فارغ
                g_zwork = st.text_input("وظيفته في التكليف (ZWORK)", value="", key="g_zwork_input")
                g_zloc = st.text_input("مكان التكليف (ZLOC)", key="g_zloc_input")
                g_zcity = st.text_input("مكان السكن (ZCITY)", key="g_zcity_input")
                g_zdate = st.date_input("📅 تاريخ التكليف:", value=datetime.now(), key="g_zdate_input")
                submit_guard = st.form_submit_button("💾 إضافة وحفظ", type="primary")
                
                if submit_guard:
                    if g_zid and g_zname:
                        # ✅ تمرير zjob2 كمتغير منفصل (9 معاملات)
                        add_other_assignment('guards', g_zid, g_zname, g_zjob, g_ZSCHOOL, g_zwork, g_zloc, g_zcity, g_zdate.strftime("%Y/%m/%d"))
                        st.success("✅ تم الإضافة بنجاح!")
                        st.session_state.guard_form_clear = True
                        st.rerun()
                    else:
                        st.error("⚠️ يرجى إدخال الهوية والاسم")
        
        if st.session_state.guard_form_clear:
            st.session_state.guard_form_clear = False
        
        with col2:
            st.markdown("**📋 قائمة الحرس**")
            df_guard = get_other_assignments('guards')
            if not df_guard.empty:
                st.dataframe(df_guard, use_container_width=True)
                
                col_btn1, col_btn2 = st.columns(2)
                with col_btn1:
                    if st.button("📦 إنشاء كتب Word للجميع", type="primary", key="btn_word_guard"):
                        docs = []
                        for _, row in df_guard.iterrows():
                            doc = generate_other_letter(row)
                            if doc:
                                bio = io.BytesIO(); doc.save(bio); bio.seek(0)
                                docs.append((bio, f"تكليف_{row['zname']}_{row['zid']}.docx"))
                        if docs:
                            st.success(f"✅ تم إنشاء {len(docs)} ملف")
                            for bio, fname in docs[:5]:
                                st.download_button(f"📥 {fname}", bio.getvalue(), fname, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_guard_{fname}")
                
                with col_btn2:
                    if st.button("📊 تصدير Excel", key="btn_excel_guard"):
                        output = io.BytesIO()
                        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                            df_guard.to_excel(writer, index=False, sheet_name='الحرس')
                            wb = writer.book; ws = writer.sheets['الحرس']
                            h_fmt = wb.add_format({'font_size': 14, 'bold': True, 'align': 'center', 'valign': 'vcenter', 'bg_color': '#1a1c23', 'font_color': '#00ffcc', 'border': 1})
                            c_fmt = wb.add_format({'font_size': 14, 'bold': True, 'align': 'right', 'valign': 'vcenter', 'border': 1})
                            ws.right_to_left(); ws.set_landscape(); ws.fit_to_pages(1, 0); ws.set_default_row(height=30)
                            for cn, val in enumerate(df_guard.columns): ws.write(0, cn, val, h_fmt)
                            for rn in range(len(df_guard)):
                                for cn in range(len(df_guard.columns)): ws.write(rn+1, cn, df_guard.iloc[rn, cn], c_fmt)
                            for idx, col in enumerate(df_guard.columns): ws.set_column(idx, idx, min(max(df_guard[col].astype(str).map(len).max(), len(str(col)))+4, 30), c_fmt)
                        st.download_button("📥 تحميل Excel", output.getvalue(), "الحرس.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", key="dl_guard_excel")
                
                st.divider()
                # ✅ تحسين الحذف: قائمة منسدلة لاختيار السجل
                df_guard_display = df_guard[['id', 'zname', 'zjob', 'ZSCHOOL']].copy()
                df_guard_display.columns = ['ID', 'الاسم', 'المهمة', 'الوظيفة الحالية']
                delete_choice = st.selectbox("🗑️ اختر السجل للحذف:", 
                                            [""] + [f"ID:{r['ID']} | {r['الاسم']} | {r['المهمة']}" for _, r in df_guard_display.iterrows()],
                                            key="guard_delete_select")
                if st.button("🗑️ حذف السجل المحدد", key="btn_del_guard"):  # ✅ أضف key هنا
                    if delete_choice and delete_choice != "":
                        selected_id = int(delete_choice.split('|')[0].replace('ID:', '').strip())
                        delete_other_assignment('guards', selected_id)
                        st.success("✅ تم الحذف")
                        st.rerun()
            else:
                st.info("📭 لا يوجد تكليفات حتى الآن")
         
    
        # ==================== تبويب مرافقة الطرود ====================
    with tab_parcels:
        st.markdown("### 📦 إدارة تكليفات مرافقة الطرود")
        
        if 'parcels_form_clear' not in st.session_state:
            st.session_state.parcels_form_clear = False
        
        col1, col2 = st.columns([1, 2])
        with col1:
            st.markdown("**➕ إضافة تكليف جديد**")
            with st.form("add_parcels_form"):
                p_zid = st.text_input("رقم الهوية (ZID)", key="p_zid_input")
                p_zname = st.text_input("الاسم (ZNAME)", key="p_zname_input")
                p_zjob = st.text_input("المهمة (ZJOB)", value="مرافق طرود", key="p_zjob_input")
                p_ZSCHOOL = st.text_input("الوظيفة الحالية (ZSCHOOL)", value="", key="p_ZSCHOOL_input")  # ✅ p_ بدلاً من g_
                p_zwork = st.text_input("وظيفته في التكليف (ZWORK)", value="", key="p_zwork_input")
                p_zloc = st.text_input("مكان التكليف (ZLOC)", key="p_zloc_input")
                p_zcity = st.text_input("مكان السكن (ZCITY)", key="p_zcity_input")
                p_zdate = st.date_input("📅 تاريخ التكليف:", value=datetime.now(), key="p_zdate_input")
                submit_parcels = st.form_submit_button("💾 إضافة وحفظ", type="primary")
                
                if submit_parcels:
                    if p_zid and p_zname:
                        # ✅ تمرير zjob2 كمتغير منفصل
                        add_other_assignment('parcels', p_zid, p_zname, p_zjob, p_ZSCHOOL, p_zwork, p_zloc, p_zcity, p_zdate.strftime("%Y/%m/%d"))
                        st.success("✅ تم الإضافة بنجاح!")
                        st.session_state.parcels_form_clear = True
                        st.rerun()
                    else:
                        st.error("⚠️ يرجى إدخال الهوية والاسم")
        
        if st.session_state.parcels_form_clear:
            st.session_state.parcels_form_clear = False
        
        with col2:
            st.markdown("**📋 قائمة مرافقة الطرود**")
            df_parcels = get_other_assignments('parcels')
            if not df_parcels.empty:
                st.dataframe(df_parcels, use_container_width=True)
                
                col_btn1, col_btn2 = st.columns(2)
                with col_btn1:
                    if st.button("📦 إنشاء كتب Word للجميع", type="primary", key="btn_word_parcels"):
                        docs = []
                        for _, row in df_parcels.iterrows():
                            doc = generate_other_letter(row)
                            if doc:
                                bio = io.BytesIO(); doc.save(bio); bio.seek(0)
                                docs.append((bio, f"تكليف_{row['zname']}_{row['zid']}.docx"))
                        if docs:
                            st.success(f"✅ تم إنشاء {len(docs)} ملف")
                            for bio, fname in docs[:5]:
                                st.download_button(f"📥 {fname}", bio.getvalue(), fname, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_parcels_{fname}")
                
                with col_btn2:
                    if st.button("📊 تصدير Excel", key="btn_excel_parcels"):
                        output = io.BytesIO()
                        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                            df_parcels.to_excel(writer, index=False, sheet_name='مرافقة_الطرود')
                            wb = writer.book; ws = writer.sheets['مرافقة_الطرود']
                            h_fmt = wb.add_format({'font_size': 14, 'bold': True, 'align': 'center', 'valign': 'vcenter', 'bg_color': '#1a1c23', 'font_color': '#00ffcc', 'border': 1})
                            c_fmt = wb.add_format({'font_size': 14, 'bold': True, 'align': 'right', 'valign': 'vcenter', 'border': 1})
                            ws.right_to_left(); ws.set_landscape(); ws.fit_to_pages(1, 0); ws.set_default_row(height=30)
                            for cn, val in enumerate(df_parcels.columns): ws.write(0, cn, val, h_fmt)
                            for rn in range(len(df_parcels)):
                                for cn in range(len(df_parcels.columns)): ws.write(rn+1, cn, df_parcels.iloc[rn, cn], c_fmt)
                            for idx, col in enumerate(df_parcels.columns): ws.set_column(idx, idx, min(max(df_parcels[col].astype(str).map(len).max(), len(str(col)))+4, 30), c_fmt)
                        st.download_button("📥 تحميل Excel", output.getvalue(), "مرافقة_الطرود.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", key="dl_parcels_excel")
                
                st.divider()
                # ✅ تحسين الحذف: قائمة منسدلة لاختيار السجل
                df_parcels_display = df_parcels[['id', 'zname', 'zjob', 'ZSCHOOL']].copy()
                df_parcels_display.columns = ['ID', 'الاسم', 'المهمة', 'الوظيفة الحالية']
                delete_choice = st.selectbox("🗑️ اختر السجل للحذف:", 
                                            [""] + [f"ID:{r['ID']} | {r['الاسم']} | {r['المهمة']}" for _, r in df_parcels_display.iterrows()],
                                            key="parcels_delete_select")
                if st.button("🗑️ حذف السجل المحدد", key="btn_del_parcels"):
                    if delete_choice and delete_choice != "":
                        selected_id = int(delete_choice.split('|')[0].replace('ID:', '').strip())
                        delete_other_assignment('parcels', selected_id)
                        st.success("✅ تم الحذف"); st.rerun()
            else:
                st.info("📭 لا يوجد تكليفات حتى الآن")
    
        # ==================== تبويب جهاز الامتحان ====================
    with tab_device:
        st.markdown("### 📱 إدارة تكليفات جهاز الامتحان")
        
        if 'device_form_clear' not in st.session_state:
            st.session_state.device_form_clear = False
        
        col1, col2 = st.columns([1, 2])
        with col1:
            st.markdown("**➕ إضافة تكليف جديد**")
            with st.form("add_device_form"):
                d_zid = st.text_input("رقم الهوية (ZID)", key="d_zid_input")
                d_zname = st.text_input("الاسم (ZNAME)", key="d_zname_input")
                d_zjob = st.text_input("المهمة (ZJOB)", value="جهاز امتحان", key="d_zjob_input")
                d_ZSCHOOL = st.text_input("الوظيفة الحالية (ZSCHOOL)", value="", key="d_ZSCHOOL_input")  # ✅ d_ بدلاً من g_
                d_zwork = st.text_input("وظيفته في التكليف (ZWORK)", value="", key="d_zwork_input")
                d_zloc = st.text_input("مكان التكليف (ZLOC)", key="d_zloc_input")
                d_zcity = st.text_input("مكان السكن (ZCITY)", key="d_zcity_input")
                d_zdate = st.date_input("📅 تاريخ التكليف:", value=datetime.now(), key="d_zdate_input")
                submit_device = st.form_submit_button("💾 إضافة وحفظ", type="primary")
                
                if submit_device:
                    if d_zid and d_zname:
                        # ✅ تمرير zjob2 كمتغير منفصل
                        add_other_assignment('exam_device', d_zid, d_zname, d_zjob, d_ZSCHOOL, d_zwork, d_zloc, d_zcity, d_zdate.strftime("%Y/%m/%d"))
                        st.success("✅ تم الإضافة بنجاح!")
                        st.session_state.device_form_clear = True
                        st.rerun()
                    else:
                        st.error("⚠️ يرجى إدخال الهوية والاسم")
        
        if st.session_state.device_form_clear:
            st.session_state.device_form_clear = False
        
        with col2:
            st.markdown("**📋 قائمة جهاز الامتحان**")
            df_device = get_other_assignments('exam_device')
            if not df_device.empty:
                st.dataframe(df_device, use_container_width=True)
                
                col_btn1, col_btn2 = st.columns(2)
                with col_btn1:
                    if st.button("📦 إنشاء كتب Word للجميع", type="primary", key="btn_word_device"):
                        docs = []
                        for _, row in df_device.iterrows():
                            doc = generate_other_letter(row)
                            if doc:
                                bio = io.BytesIO(); doc.save(bio); bio.seek(0)
                                docs.append((bio, f"تكليف_{row['zname']}_{row['zid']}.docx"))
                        if docs:
                            st.success(f"✅ تم إنشاء {len(docs)} ملف")
                            for bio, fname in docs[:5]:
                                st.download_button(f"📥 {fname}", bio.getvalue(), fname, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_device_{fname}")
                
                with col_btn2:
                   if st.button("📊 تصدير Excel", key="btn_excel_device"):
                        output = io.BytesIO()
                        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                            df_device.to_excel(writer, index=False, sheet_name='جهاز_الامتحان')
                            wb = writer.book; ws = writer.sheets['جهاز_الامتحان']
                            h_fmt = wb.add_format({'font_size': 14, 'bold': True, 'align': 'center', 'valign': 'vcenter', 'bg_color': '#1a1c23', 'font_color': '#00ffcc', 'border': 1})
                            c_fmt = wb.add_format({'font_size': 14, 'bold': True, 'align': 'right', 'valign': 'vcenter', 'border': 1})
                            ws.right_to_left(); ws.set_landscape(); ws.fit_to_pages(1, 0); ws.set_default_row(height=30)
                            for cn, val in enumerate(df_device.columns): ws.write(0, cn, val, h_fmt)
                            for rn in range(len(df_device)):
                                for cn in range(len(df_device.columns)): ws.write(rn+1, cn, df_device.iloc[rn, cn], c_fmt)
                            for idx, col in enumerate(df_device.columns): ws.set_column(idx, idx, min(max(df_device[col].astype(str).map(len).max(), len(str(col)))+4, 30), c_fmt)
                        st.download_button("📥 تحميل Excel", output.getvalue(), "جهاز_الامتحان.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", key="dl_device_excel")
                
                st.divider()
                # ✅ تحسين الحذف: قائمة منسدلة لاختيار السجل
                df_device_display = df_device[['id', 'zname', 'zjob', 'ZSCHOOL']].copy()
                df_device_display.columns = ['ID', 'الاسم', 'المهمة', 'الوظيفة الحالية']
                delete_choice = st.selectbox("🗑️ اختر السجل للحذف:", 
                                            [""] + [f"ID:{r['ID']} | {r['الاسم']} | {r['المهمة']}" for _, r in df_device_display.iterrows()],
                                            key="device_delete_select")
                if st.button("🗑️ حذف السجل المحدد", key="btn_del_device"):
                    if delete_choice and delete_choice != "":
                        selected_id = int(delete_choice.split('|')[0].replace('ID:', '').strip())
                        delete_other_assignment('exam_device', selected_id)
                        st.success("✅ تم الحذف"); st.rerun()
            else:
                st.info("📭 لا يوجد تكليفات حتى الآن")
    
        # ==================== تبويب لجنة الامتحان ====================
    with tab_committee:
        st.markdown("### 👥 إدارة تكليفات لجنة الامتحان")
        
        if 'committee_form_clear' not in st.session_state:
            st.session_state.committee_form_clear = False
        
        col1, col2 = st.columns([1, 2])
        with col1:
            st.markdown("**➕ إضافة تكليف جديد**")
            with st.form("add_committee_form"):
                c_zid = st.text_input("رقم الهوية (ZID)", key="c_zid_input")
                c_zname = st.text_input("الاسم (ZNAME)", key="c_zname_input")
                c_zjob = st.text_input("المهمة (ZJOB)", value="عضو لجنة امتحان", key="c_zjob_input")
                c_ZSCHOOL = st.text_input("الوظيفة الحالية (ZSCHOOL)", value="", key="c_ZSCHOOL_input")  # ✅ c_ بدلاً من g_
                c_zwork = st.text_input("وظيفته في التكليف (ZWORK)", value="", key="c_zwork_input")
                c_zloc = st.text_input("مكان التكليف (ZLOC)", key="c_zloc_input")
                c_zcity = st.text_input("مكان السكن (ZCITY)", key="c_zcity_input")
                c_zdate = st.date_input("📅 تاريخ التكليف:", value=datetime.now(), key="c_zdate_input")
                submit_committee = st.form_submit_button("💾 إضافة وحفظ", type="primary")
                
                if submit_committee:
                    if c_zid and c_zname:
                        # ✅ تمرير zjob2 كمتغير منفصل
                        add_other_assignment('exam_committee', c_zid, c_zname, c_zjob, c_ZSCHOOL, c_zwork, c_zloc, c_zcity, c_zdate.strftime("%Y/%m/%d"))
                        st.success("✅ تم الإضافة بنجاح!")
                        st.session_state.committee_form_clear = True
                        st.rerun()
                    else:
                        st.error("⚠️ يرجى إدخال الهوية والاسم")
        
        if st.session_state.committee_form_clear:
            st.session_state.committee_form_clear = False
        
        with col2:
            st.markdown("**📋 قائمة لجنة الامتحان**")
            df_committee = get_other_assignments('exam_committee')
            if not df_committee.empty:
                st.dataframe(df_committee, use_container_width=True)
    
                col_btn1, col_btn2 = st.columns(2)
                with col_btn1:
                    if st.button("📦 إنشاء كتب Word للجميع", type="primary", key="btn_word_committee"):
                        docs = []
                        for _, row in df_committee.iterrows():
                            doc = generate_other_letter(row)
                            if doc:
                                bio = io.BytesIO(); doc.save(bio); bio.seek(0)
                                docs.append((bio, f"تكليف_{row['zname']}_{row['zid']}.docx"))
                        if docs:
                            st.success(f"✅ تم إنشاء {len(docs)} ملف")
                            for bio, fname in docs[:5]:
                                st.download_button(f"📥 {fname}", bio.getvalue(), fname, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_committee_{fname}")
                
                with col_btn2:
                    if st.button("📊 تصدير Excel", key="btn_excel_committee"):
                        output = io.BytesIO()
                        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                            df_committee.to_excel(writer, index=False, sheet_name='لجنة_الامتحان')
                            wb = writer.book; ws = writer.sheets['لجنة_الامتحان']
                            h_fmt = wb.add_format({'font_size': 14, 'bold': True, 'align': 'center', 'valign': 'vcenter', 'bg_color': '#1a1c23', 'font_color': '#00ffcc', 'border': 1})
                            c_fmt = wb.add_format({'font_size': 14, 'bold': True, 'align': 'right', 'valign': 'vcenter', 'border': 1})
                            ws.right_to_left(); ws.set_landscape(); ws.fit_to_pages(1, 0); ws.set_default_row(height=30)
                            for cn, val in enumerate(df_committee.columns): ws.write(0, cn, val, h_fmt)
                            for rn in range(len(df_committee)):
                                for cn in range(len(df_committee.columns)): ws.write(rn+1, cn, df_committee.iloc[rn, cn], c_fmt)
                            for idx, col in enumerate(df_committee.columns): ws.set_column(idx, idx, min(max(df_committee[col].astype(str).map(len).max(), len(str(col)))+4, 30), c_fmt)
                        st.download_button("📥 تحميل Excel", output.getvalue(), "لجنة_الامتحان.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", key="dl_committee_excel")
                
                st.divider()
                # ✅ تحسين الحذف: قائمة منسدلة لاختيار السجل
                df_committee_display = df_committee[['id', 'zname', 'zjob', 'ZSCHOOL']].copy()
                df_committee_display.columns = ['ID', 'الاسم', 'المهمة', 'الوظيفة الحالية']
                delete_choice = st.selectbox("🗑️ اختر السجل للحذف:", 
                                            [""] + [f"ID:{r['ID']} | {r['الاسم']} | {r['المهمة']}" for _, r in df_committee_display.iterrows()],
                                            key="committee_delete_select")
                if st.button("🗑️ حذف السجل المحدد", key="btn_del_committee"):
                    if delete_choice and delete_choice != "":
                        selected_id = int(delete_choice.split('|')[0].replace('ID:', '').strip())
                        delete_other_assignment('exam_committee', selected_id)
                        st.success("✅ تم الحذف"); st.rerun()
            else:
                st.info("📭 لا يوجد تكليفات حتى الآن")
    
    st.divider()
    st.info("📌 **ملاحظة:** الرموز المطلوبة في القالب: ZID, ZNAME, ZJOB, ZSCHOOL, ZWORK, ZLOC, ZCITY, ZDATE")
    
    conn_other.close()
    st.stop()
